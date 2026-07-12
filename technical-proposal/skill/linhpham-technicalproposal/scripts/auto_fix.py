#!/usr/bin/env python3
"""Apply auto-fixable patches to the assembled .docx based on a format_review.json issue id.

Supported issue ids:
  - track_changes_present     -> accept all revisions
  - comments_present          -> remove all comments
  - body_not_justified        -> set every non-heading paragraph to JUSTIFY
  - unfilled_placeholder      -> remove literal {{TOKEN}} occurrences (last-resort)

Each fix mutates the .docx in place.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn


def fix_track_changes(docx_path: Path) -> int:
    # Strip <w:ins>/<w:del>/<w:moveFrom>/<w:moveTo> wrappers, keeping inserted text and dropping deleted.
    import shutil, tempfile
    tmp = docx_path.with_suffix(".fix.docx")
    fixed = 0
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in ("word/document.xml", "word/footnotes.xml", "word/endnotes.xml"):
                txt = data.decode("utf-8", errors="replace")
                # Remove deletions outright
                new = re.sub(r"<w:del\b[^>]*>.*?</w:del>", "", txt, flags=re.DOTALL)
                new = re.sub(r"<w:moveFrom\b[^>]*>.*?</w:moveFrom>", "", new, flags=re.DOTALL)
                # Unwrap insertions / moveTo (keep inner content)
                new = re.sub(r"<w:ins\b[^>]*>(.*?)</w:ins>", r"\1", new, flags=re.DOTALL)
                new = re.sub(r"<w:moveTo\b[^>]*>(.*?)</w:moveTo>", r"\1", new, flags=re.DOTALL)
                if new != txt:
                    fixed += 1
                data = new.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, docx_path)
    return fixed


def fix_comments(docx_path: Path) -> int:
    import shutil
    tmp = docx_path.with_suffix(".fix.docx")
    fixed = 0
    skip_files = {"word/comments.xml", "word/commentsExtended.xml", "word/commentsIds.xml",
                  "word/commentsExtensible.xml"}
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in skip_files:
                fixed += 1
                continue
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                txt = data.decode("utf-8", errors="replace")
                # Remove comment references and ranges
                new = re.sub(r"<w:commentRangeStart\b[^>]*/>", "", txt)
                new = re.sub(r"<w:commentRangeEnd\b[^>]*/>", "", new)
                new = re.sub(r"<w:commentReference\b[^>]*/>", "", new)
                data = new.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, docx_path)
    return fixed


def fix_justify(docx_path: Path) -> int:
    doc = Document(str(docx_path))
    count = 0
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if sname.startswith("Heading") or sname in ("Title", "Subtitle", "Caption"):
            continue
        if not p.text.strip():
            continue
        if p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            count += 1
    doc.save(str(docx_path))
    return count


def fix_placeholders(docx_path: Path) -> int:
    doc = Document(str(docx_path))
    count = 0
    for p in doc.paragraphs:
        if "{{" in p.text and "}}" in p.text:
            new = re.sub(r"\{\{[A-Z_][A-Z0-9_]*\}\}", "", p.text)
            runs = p.runs
            if runs:
                runs[0].text = new
                for r in runs[1:]:
                    r.text = ""
            count += 1
    doc.save(str(docx_path))
    return count


_MD_EMPH = r"(?:\*\*|__|\*|_)"
_LEADING_NUMBER_RE = re.compile(
    r"^\s*●\s*"
    + _MD_EMPH + r"?"                          # tolerate **/__ wrappers
    + r"\s*"
    + r"(?:"
    + r"\(?\d{1,3}\)?\s*[.)\-:]\s+"
    + r"|step\s+\d{1,3}[a-z]?\s*[.)\-:—]\s+"   # "Step 2a — " too
    + r")",
    re.IGNORECASE,
)


def fix_bullet_number_duplicate(docx_path: Path) -> int:
    """Remove ALL manual numbering between the bullet glyph and the rest of
    a bullet item. Loops per paragraph because the content-writer agent
    sometimes stacks two numberings (`●  1. Step 1 — text`) — one pass
    strips "1.", a second pass strips "Step 1 —". Preserves paragraph
    styling by writing through runs."""
    doc = Document(str(docx_path))
    count = 0
    for p in doc.paragraphs:
        joined = "".join(r.text or "" for r in p.runs)
        if not joined.lstrip().startswith("●"):
            continue
        rest = joined.lstrip()
        changed = False
        # Strip up to 3 stacked numberings; cap so we never infinite-loop on
        # pathological input.
        for _ in range(3):
            m = _LEADING_NUMBER_RE.match(rest)
            if not m:
                break
            # Keep "●  " sentinel; drop everything else matched.
            rest = "●  " + rest[m.end():]
            changed = True
        if not changed:
            continue
        runs = p.runs
        if not runs:
            continue
        runs[0].text = rest
        for r in runs[1:]:
            r.text = ""
        count += 1
    doc.save(str(docx_path))
    return count


def fix_image_text_crush(docx_path: Path) -> int:
    """Ensure every image paragraph has space-before >= 12pt and space-after >= 6pt."""
    doc = Document(str(docx_path))
    count = 0
    for p in doc.paragraphs:
        has_image = any(
            r.element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")
            for r in p.runs
        )
        if not has_image:
            continue
        pf = p.paragraph_format
        sb = pf.space_before.pt if pf.space_before else 0
        sa = pf.space_after.pt if pf.space_after else 0
        changed = False
        if sb < 12:
            pf.space_before = Pt(12)
            changed = True
        if sa < 6:
            pf.space_after = Pt(6)
            changed = True
        if changed:
            count += 1
    doc.save(str(docx_path))
    return count


def fix_heading_deeper_than_body(docx_path: Path) -> int:
    """Zero out any positive left-indent on headings AND zero a hanging
    first_line_indent if the resulting position would push the auto-number
    off the left page margin (the "2.1.x tràn ra ngoài" defect)."""
    doc = Document(str(docx_path))
    count = 0
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if not sname.startswith("Heading"):
            continue
        pf = p.paragraph_format
        if pf.left_indent is None or pf.left_indent.pt <= 0:
            continue
        pf.left_indent = Pt(0)
        fl_pt = pf.first_line_indent.pt if pf.first_line_indent else 0
        if fl_pt < 0:
            pf.first_line_indent = Pt(0)
        count += 1
    doc.save(str(docx_path))
    return count


def fix_heading_section_spacing(docx_path: Path) -> int:
    """Raise heading space-before/after to the documented floors (kept in sync
    with build_docx `_HEADING_MIN_*`). A heading stacked directly under another
    heading is skipped — its small space-before is intentional (no blank-line
    gap between "1." and "1.1")."""
    floors_before = {1: 18, 2: 12, 3: 10, 4: 8, 5: 8, 6: 6}
    floors_after  = {1:  6, 2:  4, 3:  4, 4: 3, 5: 3, 6: 3}
    doc = Document(str(docx_path))
    heading_ids = {s.style_id for s in doc.styles
                   if getattr(s, "style_id", None) and (s.name or "").startswith("Heading")}

    def prev_is_heading(p_el):
        def is_h(style):
            return style in heading_ids or style.startswith("Heading")
        prev = p_el.getprevious()
        while prev is not None:
            if prev.tag != qn("w:p"):
                return False
            pPr = prev.find(qn("w:pPr"))
            style = ""
            if pPr is not None:
                ps = pPr.find(qn("w:pStyle"))
                if ps is not None:
                    style = ps.get(qn("w:val")) or ""
            text = "".join(t.text or "" for t in prev.iter(qn("w:t")))
            if not text.strip() and not is_h(style):
                prev = prev.getprevious()
                continue
            return is_h(style)
        return False
    count = 0
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if not sname.startswith("Heading"):
            continue
        try:
            lvl = int(sname.split()[-1])
        except ValueError:
            continue
        sb_floor = floors_before.get(lvl)
        sa_floor = floors_after.get(lvl)
        if sb_floor is None:
            continue
        pf = p.paragraph_format
        changed = False
        if not prev_is_heading(p._p):
            cur_sb = pf.space_before.pt if pf.space_before else 0
            if cur_sb + 0.5 < sb_floor:
                pf.space_before = Pt(sb_floor)
                changed = True
        cur_sa = pf.space_after.pt if pf.space_after else 0
        if cur_sa + 0.5 < sa_floor:
            pf.space_after = Pt(sa_floor)
            changed = True
        if changed:
            count += 1
    doc.save(str(docx_path))
    return count


def fix_settings_flags(docx_path: Path) -> int:
    """Insert the document-level flags into settings.xml."""
    import shutil
    tmp = docx_path.with_suffix(".fix.docx")
    desired = {
        "autoHyphenation":            '<w:autoHyphenation w:val="true"/>',
        "consecutiveHyphenLimit":     '<w:consecutiveHyphenLimit w:val="2"/>',
        "hyphenationZone":            '<w:hyphenationZone w:val="288"/>',
        "doNotExpandShiftReturn":     '<w:doNotExpandShiftReturn/>',
        "characterSpacingControl":    '<w:characterSpacingControl w:val="compressPunctuation"/>',
        "updateFields":               '<w:updateFields w:val="true"/>',
    }
    inserted = 0
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                txt = data.decode("utf-8", errors="replace")
                add = [xml for tag, xml in desired.items() if tag not in txt]
                if add:
                    txt = txt.replace("</w:settings>", "".join(add) + "</w:settings>")
                    inserted = len(add)
                data = txt.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, docx_path)
    return inserted


def fix_caption_text_crush(docx_path: Path) -> int:
    """For every 1x1 Figure caption table, set the next paragraph's
    space-before >= 12pt. Same root fix as ensure_caption_followed_by_gap
    in build_docx but operates on an already-assembled docx."""
    from docx.oxml import OxmlElement
    doc = Document(str(docx_path))
    count = 0
    for tbl in doc.tables:
        rows = tbl.rows
        if len(rows) != 1 or len(rows[0].cells) != 1:
            continue
        cell_text = "".join(p.text for p in rows[0].cells[0].paragraphs).strip()
        if not cell_text or ("Figure" not in cell_text and "figure" not in cell_text
                              and "Bảng" not in cell_text):
            continue
        nxt = tbl._element.getnext()
        if nxt is None or nxt.tag.rsplit("}", 1)[-1] != "p":
            continue
        pPr = nxt.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            nxt.insert(0, pPr)
        sp = pPr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            pPr.append(sp)
        v = sp.get(qn("w:before"))
        cur_pt = int(v) / 20 if v else 0
        if cur_pt < 12:
            sp.set(qn("w:before"), str(12 * 20))
            count += 1
    doc.save(str(docx_path))
    return count


FIXERS = {
    "track_changes_present":          fix_track_changes,
    "comments_present":               fix_comments,
    "body_not_justified":             fix_justify,
    "unfilled_placeholder":           fix_placeholders,
    "bullet_number_duplicate":        fix_bullet_number_duplicate,
    "image_text_crush":               fix_image_text_crush,
    "heading_deeper_than_body":       fix_heading_deeper_than_body,
    "heading_section_spacing_tight":  fix_heading_section_spacing,
    "caption_text_crush":             fix_caption_text_crush,
    "settings_flags_missing":         fix_settings_flags,
    "justify_soft_returns":           fix_settings_flags,
}


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", type=Path, required=True)
    ap.add_argument("--issue", required=True, help="Issue id from format_review.json")
    args = ap.parse_args()

    if not args.docx.exists():
        print(f"! docx not found: {args.docx}", file=sys.stderr); return 1

    issue = args.issue.split("::", 1)[0]

    fixer = FIXERS.get(issue)
    if not fixer:
        print(f"! no auto-fix registered for issue: {issue}", file=sys.stderr); return 1

    n = fixer(args.docx)
    print(f"{issue}: {n} item(s) patched")
    return 0


if __name__ == "__main__":
    sys.exit(main())
