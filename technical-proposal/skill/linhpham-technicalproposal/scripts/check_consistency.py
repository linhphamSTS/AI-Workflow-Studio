#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Does the technology stack agree with the architecture it is supposed to implement?

    python check_consistency.py --plan spec/plan.json \\
                                --diagrams output/diagrams/diagrams.json \\
                                --docx "output/<Project>.docx"

The stack, the architecture prose and the figures are produced by different steps, and
nothing compared them. They can therefore disagree, and they disagree silently: every
individual artefact looks correct on its own. Two real ways it goes wrong.

**Two clouds in one bid.** The stack names one provider, a figure is drawn from another
provider's template, and a reader concludes the bidder does not know what they are
proposing. It survives review because no single document contains the contradiction.

**A figure promising a capability the stack cannot deliver.** A diagram box reading
"Kafka protocol stream" and "Schema registry" beside a stack row naming a plain message
queue is a commitment with nothing behind it: no replayable log, no schema governance. The
figure is the thing an evaluator believes, so the stack has to be able to honour it.

Exit 0 only when the three agree.
"""
import argparse
import json
import os
import re
import sys

# ---------------------------------------------------------------- cloud vocabulary
# Only names that identify a provider beyond doubt. Ambiguous words ("compute", "storage",
# "functions") are deliberately absent: a vocabulary that catches a generic noun produces a
# false contradiction, and a gate that cries wolf gets switched off.
CLOUDS = {
    'AWS': (r'\bAWS\b', r'\bAmazon Web Services\b', r'\bEKS\b', r'\bECS\b', r'\bFargate\b',
            r'\bAurora\b', r'\bDynamoDB\b', r'\bS3\b', r'\bCloudFront\b', r'\bRoute 53\b',
            r'\bLambda\b', r'\bEC2\b', r'\bRDS\b', r'\bElastiCache\b', r'\bOpenSearch\b',
            r'\bEventBridge\b', r'\bSQS\b', r'\bSNS\b', r'\bMSK\b', r'\bBedrock\b',
            r'\bCloudWatch\b', r'\bme-central-1\b'),
    'Azure': (r'\bAzure\b', r'\bAKS\b', r'\bCosmos DB\b', r'\bBlob Storage\b',
              r'\bApplication Gateway\b', r'\bAPI Management\b', r'\bEvent Hubs\b',
              r'\bService Bus\b', r'\bKey Vault\b', r'\bLog Analytics\b',
              r'\bFoundry Models\b', r'\buaenorth\b', r'\bUAE North\b'),
    'GCP': (r'\bGCP\b', r'\bGoogle Cloud\b', r'\bBigQuery\b', r'\bGKE\b', r'\bSpanner\b',
            r'\bPub/Sub\b', r'\bCloud Run\b', r'\bFirestore\b', r'\bVertex AI\b'),
}

# A figure that promises the left-hand capability needs the stack to name something on the
# right. This is the class of defect where the picture commits and the stack cannot pay.
PROMISES = (
    ('a replayable event log',
     (r'kafka', r'\breplay\b', r'\boffset\b', r'event stream'),
     (r'kafka', r'event hubs', r'\bmsk\b', r'kinesis', r'pulsar', r'redpanda')),
    ('schema governance on the event bus',
     (r'schema registry', r'schema governance'),
     (r'schema registry', r'kafka', r'event hubs', r'confluent', r'apicurio')),
    ('a queue with dead-letter handling',
     (r'dead.letter', r'\bdlq\b'),
     (r'queue', r'service bus', r'\bsqs\b', r'rabbit', r'kafka', r'event hubs')),
    ('a search tier with language analysers',
     (r'analyser', r'analyzer', r'full.text search', r'search index'),
     (r'search', r'opensearch', r'elastic', r'solr', r'ai search', r'algolia')),
    ('a cache tier',
     (r'\bcache\b', r'\bcaching tier\b'),
     (r'redis', r'valkey', r'memcached', r'elasticache', r'\bcache\b')),
    ('a relational store',
     (r'\bacid\b', r'double.entry', r'\bledger\b', r'transactional integrity'),
     (r'postgres', r'postgresql', r'mysql', r'sql server', r'aurora', r'oracle',
      r'mariadb', r'cockroach')),
    ('an object store',
     (r'object storage', r'blob storage', r'\bbucket\b'),
     (r'\bs3\b', r'blob storage', r'cloud storage', r'object storage', r'minio')),
)

problems, warnings, checks = [], [], []


def ck(label, ok, detail=''):
    checks.append((label, bool(ok)))
    if not ok:
        problems.append('%s: %s' % (label, detail) if detail else label)


def warn(label, detail=''):
    warnings.append('%s: %s' % (label, detail) if detail else label)


def clouds_in(text):
    """Which providers this text names, and the first phrase that named each."""
    out = {}
    for name, pats in CLOUDS.items():
        for p in pats:
            m = re.search(p, text, re.I)
            if m:
                out[name] = m.group(0)
                break
    return out


def flatten(obj, out=None):
    out = [] if out is None else out
    if isinstance(obj, str):
        out.append(obj)
    elif isinstance(obj, list):
        for v in obj:
            flatten(v, out)
    elif isinstance(obj, dict):
        for v in obj.values():
            flatten(v, out)
    return out


def docx_text(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts += [p.text for p in c.paragraphs]
    return '\n'.join(parts)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--plan', required=True)
    ap.add_argument('--diagrams')
    ap.add_argument('--docx')
    ap.add_argument('--allow-multi-cloud', action='store_true',
                    help='only for a bid that genuinely proposes more than one provider, '
                         'which is rare and should be a stated decision')
    a = ap.parse_args()

    plan = json.load(open(a.plan, encoding='utf-8'))
    stack_rows = plan.get('tech_stack') or []
    if isinstance(stack_rows, dict):
        stack_rows = [{'layer': k, 'choice': v} for k, v in stack_rows.items()]
    stack_text = '\n'.join(flatten(stack_rows))
    arch_text = str(plan.get('architecture') or '')

    ck('the plan states a technology stack', bool(stack_rows),
       'plan.tech_stack is empty, so there is nothing to compare the architecture against')
    ck('the plan states an architecture', bool(arch_text.strip()),
       'plan.architecture is empty')
    if not stack_rows:
        return report()

    diag_text = ''
    diag_specs = []
    if a.diagrams and os.path.exists(a.diagrams):
        dj = json.load(open(a.diagrams, encoding='utf-8'))
        diag_specs = dj if isinstance(dj, list) else (dj.get('diagrams') or [])
        diag_text = '\n'.join(flatten(dj))
    else:
        warn('no diagrams file given, so the figures were not compared',
             'pass --diagrams output/diagrams/diagrams.json')

    doc_text = ''
    if a.docx and os.path.exists(a.docx):
        doc_text = docx_text(a.docx)
    else:
        warn('no document given, so the delivered prose was not compared',
             'pass --docx once the build has run')

    # ------------------------------------------------------------------ one cloud
    sources = [('the stack', stack_text), ('the architecture', arch_text)]
    if diag_text:
        sources.append(('the figures', diag_text))
    if doc_text:
        sources.append(('the document', doc_text))

    per_source = {label: clouds_in(text) for label, text in sources}
    named = {}
    for label, found in per_source.items():
        for cloud, phrase in found.items():
            named.setdefault(cloud, []).append('%s ("%s")' % (label, phrase))

    if a.allow_multi_cloud:
        warn('multi-cloud was allowed explicitly', 'clouds named: %s' % ', '.join(sorted(named)))
    else:
        ck('exactly one cloud provider is named anywhere', len(named) <= 1,
           'this bid names %d providers: %s'
           % (len(named), '; '.join('%s via %s' % (c, ', '.join(v)) for c, v in named.items())))

    # The stack decides. A figure drawn from another provider's template is the common way
    # a second provider gets in, and the figure is what a reader looks at first.
    stack_clouds = set(per_source['the stack'])
    if stack_clouds and diag_text:
        diag_clouds = set(per_source.get('the figures') or {})
        ck('every figure uses the cloud the stack chose',
           not (diag_clouds - stack_clouds),
           'the stack names %s, the figures also name %s'
           % (', '.join(sorted(stack_clouds)) or 'none',
              ', '.join(sorted(diag_clouds - stack_clouds))))
        tmpl = {str(d.get('template') or d.get('kind') or '') for d in diag_specs
                if isinstance(d, dict)}
        wrong_tmpl = sorted(t for t in tmpl
                            for c in ('aws', 'azure', 'gcp')
                            if t.startswith(c + '_')
                            and c.upper() not in {s.upper() for s in stack_clouds}
                            and not (c == 'gcp' and 'GCP' in stack_clouds))
        ck('no figure is built from another provider template', not wrong_tmpl,
           'templates %s against a stack of %s'
           % (wrong_tmpl, ', '.join(sorted(stack_clouds))))

    # ------------------------------------------- the figures do not over-promise
    if diag_text:
        unbacked = []
        for label, promise_pats, backing_pats in PROMISES:
            promised = any(re.search(p, diag_text, re.I) for p in promise_pats)
            if not promised:
                continue
            backed = any(re.search(p, stack_text, re.I) for p in backing_pats)
            if not backed:
                unbacked.append('the figures promise %s, and the stack names nothing that '
                                'provides it' % label)
        ck('every capability a figure promises is in the stack', not unbacked,
           '; '.join(unbacked))

    # --------------------------------- the stack is actually used in the document
    if doc_text:
        absent = []
        for row in stack_rows:
            if not isinstance(row, dict):
                continue
            choice = str(row.get('choice') or row.get('name') or '')
            # Take the concrete product names out of the choice: capitalised or dotted
            # tokens. A choice written entirely in lower-case prose names no product and is
            # not something the document can be checked against.
            names = re.findall(r'\b[A-Z][A-Za-z0-9]*(?:\.[A-Za-z0-9]+)*\b', choice)
            names = [n for n in names if len(n) > 2
                     and n.lower() not in ('the', 'and', 'for', 'with', 'per', 'one', 'two',
                                           'rejected', 'trade')]
            if not names:
                continue
            if not any(re.search(r'\b%s\b' % re.escape(n), doc_text, re.I) for n in names[:4]):
                absent.append('%s (%s)' % (row.get('layer', '?'), ', '.join(names[:3])))
        # A warning, not a gate: a stack row can legitimately name an internal choice the
        # prose describes in other words, and failing on that would fire on sound proposals.
        if absent:
            warn('%d stack row(s) name a technology the document never mentions'
                 % len(absent), '; '.join(absent[:6]))

    return report()


def report():
    print('=' * 78)
    print('STACK vs ARCHITECTURE CONSISTENCY')
    print('=' * 78)
    for label, ok in checks:
        print('  [%s] %s' % ('PASS' if ok else 'FAIL', label))
    if warnings:
        print()
        for w in warnings:
            print('  [warn] %s' % w)
    print()
    if problems:
        print('RESULT: %d MISMATCH(ES)' % len(problems))
        for p in problems:
            print('  - %s' % p)
        return 1
    print('RESULT: CONSISTENT (%d checks)' % len(checks))
    return 0


if __name__ == '__main__':
    sys.exit(main())
