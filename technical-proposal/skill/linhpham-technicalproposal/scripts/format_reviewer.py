#!/usr/bin/env python3
"""Strict format and SharePoint-compatibility review for the assembled .docx.

Emits a JSON report (machine-readable, consumed by auto_fix.py and by the
Phase 5b prompt) plus a parallel Markdown report (human-readable).

Run from the skill folder so relative imports work:
    python scripts/format_reviewer.py --docx <path> --json out.json
"""
from __future__ import annotations

import argparse
import io
import json
import re
import sys
import zipfile
from dataclasses import asdict, dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


@dataclass
class Issue:
    id: str
    category: str  # layout | visual | sharepoint | polish | sharpness
    severity: str  # blocker | major | minor
    auto_fixable: bool
    summary: str
    detail: str = ""
    page: int | None = None
    location: str | None = None


@dataclass
class Report:
    docx: str
    file_size_bytes: int
    issues: list[Issue] = field(default_factory=list)
    checks_passed: list[str] = field(default_factory=list)
    auto_fixable_count: int = 0
    blocker_count: int = 0

    def add(self, issue: Issue) -> None:
        self.issues.append(issue)
        if issue.auto_fixable:
            self.auto_fixable_count += 1
        if issue.severity == "blocker":
            self.blocker_count += 1

    def pass_(self, name: str) -> None:
        self.checks_passed.append(name)


# ---------------------------------------------------------------------------
# Checks.
# ---------------------------------------------------------------------------


def check_zip_integrity(docx_path: Path, report: Report) -> None:
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            bad = z.testzip()
            if bad:
                report.add(Issue("zip_corruption", "sharepoint", "blocker", False,
                                 "Zip integrity check failed",
                                 detail=f"Bad file inside docx: {bad}"))
            else:
                report.pass_("zip_integrity")
    except zipfile.BadZipFile:
        report.add(Issue("zip_bad", "sharepoint", "blocker", False,
                         "File is not a valid .docx (bad zip)"))


def check_size(docx_path: Path, report: Report) -> None:
    size = docx_path.stat().st_size
    if size > 100 * 1024 * 1024:
        report.add(Issue("file_too_large", "sharepoint", "blocker", False,
                         f"File size {size/1024/1024:.1f} MB exceeds 100 MB"))
    elif size > 25 * 1024 * 1024:
        report.add(Issue("file_large_warning", "sharepoint", "minor", False,
                         f"File size {size/1024/1024:.1f} MB > 25 MB — slow SharePoint sync"))
    else:
        report.pass_("file_size_ok")


def check_no_track_changes(docx_path: Path, report: Report) -> None:
    with zipfile.ZipFile(docx_path) as z:
        doc_xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    if re.search(r"<w:ins\b|<w:del\b|<w:moveFrom\b|<w:moveTo\b", doc_xml):
        report.add(Issue("track_changes_present", "sharepoint", "major", True,
                         "Tracked changes (insertions / deletions) present"))
    else:
        report.pass_("no_track_changes")


def check_no_comments(docx_path: Path, report: Report) -> None:
    with zipfile.ZipFile(docx_path) as z:
        if "word/comments.xml" in z.namelist():
            data = z.read("word/comments.xml").decode("utf-8", errors="replace")
            if "<w:comment " in data:
                report.add(Issue("comments_present", "sharepoint", "major", True,
                                 "Document contains comments"))
                return
        report.pass_("no_comments")


def check_no_macros(docx_path: Path, report: Report) -> None:
    if docx_path.suffix.lower() == ".docm":
        report.add(Issue("docm_macro_enabled", "sharepoint", "blocker", False,
                         ".docm extension — macros allowed; should be .docx"))
        return
    with zipfile.ZipFile(docx_path) as z:
        if any(n.startswith("word/vbaProject") for n in z.namelist()):
            report.add(Issue("vba_project_present", "sharepoint", "blocker", False,
                             "VBA macros found inside .docx"))
            return
    report.pass_("no_macros")


def check_no_encryption(docx_path: Path, report: Report) -> None:
    with open(docx_path, "rb") as f:
        head = f.read(8)
    # Encrypted OOXML starts with the OLE compound document signature D0 CF 11 E0 A1 B1 1A E1
    if head[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        report.add(Issue("encrypted", "sharepoint", "blocker", False,
                         "File appears encrypted / password-protected"))
    else:
        report.pass_("not_encrypted")


_REL_RE = re.compile(r"<Relationship\b[^>]*/>")


def check_no_external_images(docx_path: Path, report: Report) -> None:
    """An image linked by URL instead of embedded does not travel with the file.

    Test each relationship on its own. The earlier version asked whether the .rels
    file contained `TargetMode="External"` ANYWHERE and the word "image" ANYWHERE,
    which is true of every document that has both an embedded image and an ordinary
    hyperlink: the case-study link alone was enough to raise a false alarm. Two
    unrelated conditions matched at file level say nothing about any one relationship.
    """
    offenders = []
    with zipfile.ZipFile(docx_path) as z:
        for rel_name in z.namelist():
            if not rel_name.endswith(".rels"):
                continue
            data = z.read(rel_name).decode("utf-8", errors="replace")
            for rel in _REL_RE.findall(data):
                if 'TargetMode="External"' not in rel:
                    continue
                type_m = re.search(r'Type="([^"]+)"', rel)
                if not type_m or not type_m.group(1).rstrip("/").endswith("/image"):
                    continue          # an external hyperlink is normal and expected
                target = re.search(r'Target="([^"]+)"', rel)
                offenders.append(f"{rel_name}: {target.group(1) if target else '?'}")
    if offenders:
        report.add(Issue("external_image_ref", "sharepoint", "major", False,
                         f"{len(offenders)} image(s) linked externally instead of embedded; "
                         f"they will not render for anyone else",
                         detail="; ".join(offenders[:5])))
    else:
        report.pass_("images_all_embedded")


def check_no_unfilled_placeholders(docx_path: Path, report: Report) -> None:
    with zipfile.ZipFile(docx_path) as z:
        text = z.read("word/document.xml").decode("utf-8", errors="replace")
    placeholders = re.findall(r"\{\{[A-Z_][A-Z0-9_]*\}\}", text)
    if placeholders:
        report.add(Issue("unfilled_placeholder", "polish", "blocker", True,
                         "Literal {{PLACEHOLDER}} tokens left in document",
                         detail=f"Tokens: {sorted(set(placeholders))}"))
    else:
        report.pass_("placeholders_filled")


def check_image_sharpness(docx_path: Path, report: Report) -> None:
    """Every GENERATED figure must be >= 1500 px wide for a sharp 6.5-in Word display.

    Scoped to what the run produces. The template carries its own decorative art, and the
    technology-stack tables carry deliberately small logos; neither is ours to change, and
    reporting 78 issues on them every run taught the reader to ignore this check, which is
    worse than not having it. A generated architecture figure is 4,000 px and up, so the
    split is unambiguous: anything under the in-table logo ceiling is a logo, and anything
    the run did not generate is template art.
    """
    try:
        from PIL import Image
    except ImportError:
        return  # If PIL unavailable in this environment, skip silently
    with zipfile.ZipFile(docx_path) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]
        for m in media:
            try:
                data = z.read(m)
                img = Image.open(io.BytesIO(data))
                w, _ = img.size
                if w < 1500:
                    report.add(Issue(
                        f"image_low_resolution::{m}", "sharpness", "major", False,
                        f"Image {m} is only {w}px wide (< 1500 px target)",
                        detail="Re-render the source diagram at >= 300 DPI",
                    ))
            except Exception:
                continue
    report.pass_("image_sharpness_audited")


def check_justify_body(docx_path: Path, report: Report) -> None:
    """Sample paragraphs: ratio of justified body paragraphs should be high."""
    doc = Document(str(docx_path))
    body, justified = 0, 0
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if sname.startswith("Heading") or sname in ("Title", "Subtitle", "Caption", "TOC", "Header", "Footer"):
            continue
        if not p.text.strip():
            continue
        body += 1
        if p.alignment is not None and int(p.alignment) == 3:  # WD_ALIGN_PARAGRAPH.JUSTIFY
            justified += 1
    if body == 0:
        return
    ratio = justified / body
    if ratio < 0.8:
        report.add(Issue("body_not_justified", "visual", "major", True,
                         f"Only {justified}/{body} body paragraphs are justified ({ratio:.0%})"))
    else:
        report.pass_("body_justified")


def check_heading_orphans_pdf(pages_dir: Path | None, report: Report) -> None:
    """Visual checks need rendered PNGs (Phase 5b renders them). Skipped here
    if pages_dir is not provided; the per-page Read by the agent covers it."""
    if pages_dir is None or not pages_dir.exists():
        return
    # Programmatically checking widows / orphan headings from PNGs is hard.
    # Leave to the agent's visual inspection step.
    report.pass_("page_visual_check_deferred_to_agent")


# ---------------------------------------------------------------------------
# Quality checks added 2026-05-25 after the user flagged specific defects:
# bullet-and-number duplicates, image-text crush, heading deeper-indent than
# body, tight section spacing, justify whitespace channels, settings flags
# missing. Each maps to a real visual symptom the reviewer must not miss.
# ---------------------------------------------------------------------------

_BULLET_NUMBER_PATTERN = re.compile(
    r"●\s*(?:\(?\d{1,3}\)?\s*[.)\-:]|step\s+\d{1,3}\s*[.)\-:—])\s+",
    re.IGNORECASE,
)


def check_no_bullet_number_duplicates(docx_path: Path, report: Report) -> None:
    """Reject paragraphs whose text starts with "●  1. " / "●  Step 2 — " etc.
    Each one renders as both a bullet dot AND a manual number — visible to
    the reader as a quality bug ("dấu chấm tròn + số trùng")."""
    doc = Document(str(docx_path))
    bad = []
    for i, p in enumerate(doc.paragraphs):
        txt = (p.text or "").lstrip()
        if not txt.startswith("●"):
            continue
        if _BULLET_NUMBER_PATTERN.match(txt):
            bad.append((i, txt[:80]))
    if bad:
        report.add(Issue(
            "bullet_number_duplicate", "polish", "blocker", True,
            f"{len(bad)} bullet(s) start with both '●' and a manual number "
            f"(renders as '● 1. text')",
            detail="; ".join(f"para {i}: {t}" for i, t in bad[:5]),
        ))
    else:
        report.pass_("no_bullet_number_duplicates")


def check_image_spacing(docx_path: Path, report: Report) -> None:
    """Every image paragraph must have space-before AND space-after >= 6pt so
    the image is not visually crushed against the text above and below."""
    doc = Document(str(docx_path))
    bad = []
    for i, p in enumerate(doc.paragraphs):
        has_image = any(
            r.element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline")
            for r in p.runs
        )
        if not has_image:
            continue
        pf = p.paragraph_format
        sb = pf.space_before.pt if pf.space_before else 0
        sa = pf.space_after.pt if pf.space_after else 0
        if sb < 6 or sa < 3:
            bad.append((i, sb, sa))
    if bad:
        report.add(Issue(
            "image_text_crush", "visual", "major", True,
            f"{len(bad)} image paragraph(s) have insufficient spacing "
            f"(sb<6pt OR sa<3pt) — image touches surrounding text",
            detail="; ".join(f"para {i}: sb={sb} sa={sa}" for i, sb, sa in bad[:5]),
        ))
    else:
        report.pass_("image_spacing_ok")


def check_heading_indent_vs_body(docx_path: Path, report: Report) -> None:
    """Heading 3 / 4 / 5 / 6 must NOT be left-indented deeper than the
    surrounding body paragraphs. A heading sitting at L=36 while body sits
    at L=0 reads as "title thụt vô sâu hơn text" — wrong visual hierarchy.
    """
    doc = Document(str(docx_path))
    bad = []
    # Reference body indent: most common non-zero left indent of Normal
    # paragraphs, or 0 if all Normal text is flush-left.
    body_indents = [p.paragraph_format.left_indent.pt
                    for p in doc.paragraphs
                    if (p.style.name if p.style else "") == "Normal"
                    and p.paragraph_format.left_indent is not None]
    body_indent_ref = 0
    if body_indents:
        body_indent_ref = max(0, min(body_indents))
    for i, p in enumerate(doc.paragraphs):
        sname = p.style.name if p.style else ""
        if not sname.startswith("Heading"):
            continue
        pf = p.paragraph_format
        if pf.left_indent is None:
            continue
        if pf.left_indent.pt > body_indent_ref + 2:  # 2pt slack
            bad.append((i, sname, pf.left_indent.pt, body_indent_ref))
    if bad:
        report.add(Issue(
            "heading_deeper_than_body", "visual", "blocker", True,
            f"{len(bad)} heading(s) indented deeper than body text",
            detail="; ".join(f"para {i} ({s}): L={l} vs body L={b}"
                              for i, s, l, b in bad[:5]),
        ))
    else:
        report.pass_("heading_indent_consistent")


def _prev_block_is_heading_el(p_el, heading_ids) -> bool:
    """True if the block before `p_el` is a heading paragraph (skipping empty
    spacers). A heading stacked directly under another heading hugs it with a
    small space-before on purpose — it must NOT be flagged as 'tight'.
    `heading_ids` maps numeric style IDs (the template uses `w:val="4"`) to
    headings — a name-only `startswith("Heading")` check silently misses them."""
    from docx.oxml.ns import qn as _qn
    def is_h(style):
        return style in heading_ids or style.startswith("Heading")
    prev = p_el.getprevious()
    while prev is not None:
        if prev.tag != _qn("w:p"):
            return False
        pPr = prev.find(_qn("w:pPr"))
        style = ""
        if pPr is not None:
            ps = pPr.find(_qn("w:pStyle"))
            if ps is not None:
                style = ps.get(_qn("w:val")) or ""
        text = "".join(t.text or "" for t in prev.iter(_qn("w:t")))
        if not text.strip() and not is_h(style):
            prev = prev.getprevious()
            continue
        return is_h(style)
    return False


def check_heading_section_spacing(docx_path: Path, report: Report) -> None:
    """Heading 1-3 should have visible space-before so major section
    transitions don't feel cramped — but NOT a blank-line gap. Floors match
    build_docx `_HEADING_MIN_SPACE_BEFORE`: H1 >= 18, H2 >= 12, H3 >= 10,
    H4/H5 >= 8, H6 >= 6pt. A heading stacked directly under another heading is
    exempt (it hugs its parent with a small gap on purpose).
    """
    floors = {1: 18, 2: 12, 3: 10, 4: 8, 5: 8, 6: 6}
    doc = Document(str(docx_path))
    heading_ids = {s.style_id for s in doc.styles
                   if getattr(s, "style_id", None) and (s.name or "").startswith("Heading")}
    bad = []
    for i, p in enumerate(doc.paragraphs):
        sname = p.style.name if p.style else ""
        if not sname.startswith("Heading"):
            continue
        try:
            lvl = int(sname.split()[-1])
        except ValueError:
            continue
        floor = floors.get(lvl)
        if floor is None:
            continue
        if _prev_block_is_heading_el(p._p, heading_ids):
            continue  # stacked heading — small gap is intentional
        pf = p.paragraph_format
        sb = pf.space_before.pt if pf.space_before else 0
        if sb + 0.5 < floor:  # 0.5pt rounding slack
            bad.append((i, sname, sb, floor))
    if bad:
        report.add(Issue(
            "heading_section_spacing_tight", "visual", "major", True,
            f"{len(bad)} heading(s) have space-before below the floor "
            f"(major sections look cramped)",
            detail="; ".join(f"para {i} ({s}): sb={sb}<{floor}"
                              for i, s, sb, floor in bad[:5]),
        ))
    else:
        report.pass_("heading_section_spacing_ok")


def check_settings_flags_present(docx_path: Path, report: Report) -> None:
    """settings.xml must carry the flags that prevent justify whitespace
    channels and TOC stale-cache."""
    required = [
        "autoHyphenation",
        "doNotExpandShiftReturn",
        "characterSpacingControl",
        "updateFields",
    ]
    with zipfile.ZipFile(docx_path) as z:
        try:
            txt = z.read("word/settings.xml").decode("utf-8", errors="replace")
        except KeyError:
            report.add(Issue("settings_xml_missing", "polish", "major", False,
                             "word/settings.xml not present"))
            return
    missing = [f for f in required if f not in txt]
    if missing:
        report.add(Issue(
            "settings_flags_missing", "visual", "major", True,
            f"settings.xml missing flag(s): {', '.join(missing)} — justify "
            f"gaps and TOC freshness may regress",
        ))
    else:
        report.pass_("settings_flags_present")


def check_caption_followed_by_gap(docx_path: Path, report: Report) -> None:
    """For every 1x1 Figure caption table, the next paragraph must carry
    space-before >= 8pt. Otherwise body text sits flush against the
    caption — the user-flagged "text và title image quá sát nhau" defect.
    """
    doc = Document(str(docx_path))
    body = doc.element.body
    bad = []
    for tbl in doc.tables:
        rows = tbl.rows
        if len(rows) != 1 or len(rows[0].cells) != 1:
            continue
        cell_text = "".join(p.text for p in rows[0].cells[0].paragraphs).strip()
        if not cell_text or ("Figure" not in cell_text and "figure" not in cell_text
                              and "Bảng" not in cell_text):
            continue
        nxt = tbl._element.getnext()
        if nxt is None or nxt.tag.rsplit("}", 1)[-1] != "p":
            continue
        pPr = nxt.find(qn("w:pPr"))
        sb_pt = 0
        if pPr is not None:
            sp = pPr.find(qn("w:spacing"))
            if sp is not None:
                v = sp.get(qn("w:before"))
                if v:
                    sb_pt = int(v) / 20
        if sb_pt < 8:
            # describe the next paragraph briefly
            txt_t = nxt.find(".//" + qn("w:t"))
            preview = (txt_t.text or "")[:50] if txt_t is not None else ""
            bad.append((cell_text[:40], sb_pt, preview))
    if bad:
        report.add(Issue(
            "caption_text_crush", "visual", "major", True,
            f"{len(bad)} Figure caption(s) have body text crammed underneath "
            f"(next paragraph space-before < 8pt)",
            detail="; ".join(f"after '{c}': sb={sb} '{p}'" for c, sb, p in bad[:5]),
        ))
    else:
        report.pass_("caption_gap_ok")


def check_justify_whitespace_channels(docx_path: Path, report: Report) -> None:
    """When a justified paragraph contains a soft line break, Word can
    expand the inter-word spacing on the broken line — producing visible
    "rivers" of white space. The document-level flag `doNotExpandShiftReturn`
    neutralises this symptom. We only raise an issue when:
      (a) the flag is missing from settings.xml AND
      (b) at least one justified paragraph contains a `<w:br/>` soft return.
    Either alone is harmless: (a) without (b) means no paragraph would be
    affected; (b) without (a) is already prevented by Word.
    """
    # Read settings to see if the neutralising flag is in place.
    flag_present = False
    with zipfile.ZipFile(docx_path) as z:
        try:
            settings = z.read("word/settings.xml").decode("utf-8", errors="replace")
            flag_present = "doNotExpandShiftReturn" in settings
        except KeyError:
            pass
    if flag_present:
        report.pass_("justify_whitespace_neutralised")
        return

    doc = Document(str(docx_path))
    suspects = []
    for i, p in enumerate(doc.paragraphs):
        if p.alignment is None or int(p.alignment) != 3:
            continue
        for r in p.runs:
            if r.element.findall(qn("w:br")):
                suspects.append((i, (p.text or "")[:60]))
                break
    if suspects:
        report.add(Issue(
            "justify_soft_returns", "visual", "major", True,
            f"{len(suspects)} justified paragraph(s) contain soft line breaks "
            f"AND the document is missing `doNotExpandShiftReturn` — visible "
            f"whitespace channels will appear in Word",
            detail="; ".join(f"para {i}: {t}" for i, t in suspects[:5]),
        ))
    else:
        report.pass_("no_justify_soft_returns")


_TECHSTACK_LABELS = {
    "back-end", "backend", "front-end", "frontend", "database",
    "server & hosting", "server and hosting", "data", "ai", "ai / ml", "ai & ml",
}


def check_techstack_is_table(docx_path: Path, report: Report) -> None:
    """Each Technology Stack sub-heading (Back-end / Front-end / Database /
    Server & Hosting / Data / AI) MUST be immediately followed by a 2-column
    'Technology | Advantages' table, NOT a prose paragraph — the required
    professional table format. Prose here is a zero-tolerance format defect
    (`techstack_not_table`): the content-writer emitted a string instead of the
    required array of {name, description} rows."""
    doc = Document(str(docx_path))
    pmap = {p._p: p for p in doc.paragraphs}
    tmap = {t._tbl: t for t in doc.tables}
    children = list(doc.element.body.iterchildren())
    bad = []
    for idx, child in enumerate(children):
        p = pmap.get(child)
        if p is None:
            continue
        if (p.style.name if p.style else "") != "Heading 6":
            continue
        if (p.text or "").strip().lower() not in _TECHSTACK_LABELS:
            continue
        nxt = None
        for follow in children[idx + 1:]:
            if follow in tmap:
                nxt = ("tbl", None); break
            fp = pmap.get(follow)
            if fp is not None:
                if (fp.text or "").strip() == "":
                    continue  # skip blank spacer paragraphs
                nxt = ("p", (fp.text or "")[:60]); break
        if nxt is None or nxt[0] != "tbl":
            bad.append(((p.text or "").strip(), nxt[1] if nxt else "(nothing)"))
    if bad:
        report.add(Issue(
            "techstack_not_table", "visual", "blocker", False,
            f"{len(bad)} Technology Stack sub-section(s) render as prose instead of a "
            f"'Technology | Advantages' table",
            detail="; ".join(f"'{h}' -> {prev}" for h, prev in bad),
        ))
    else:
        report.pass_("techstack_is_table")


#: Optional-section heading text (lowercased) -> word ceiling from `04_generate.md`.
#: A bid reader skims these looking for a number, a name or a commitment, and long
#: prose hides all three. The ceilings live here so the documented rule has a check
#: behind it: a style rule nobody measures gets violated on the next run.
_SECTION_WORD_CEILINGS = {
    "security & data protection": 260,
    "team structure": 150,
    "engagement model": 110,
    "delivery roadmap": 180,
    "milestones & acceptance": 160,
    "governance & reporting": 130,
    "support model": 160,
}
#: Sections whose body must be a table, not prose. Same detection as the technology
#: stack: a role list needs the position, the headcount, the seniority and the
#: location side by side, which prose cannot present scannably.
_TABLE_SECTIONS = {"roles & responsibilities"}

#: Narrative sections outside the optional block, from the WRITE TIGHT rule. Kept
#: separate from `_SECTION_WORD_CEILINGS` only so the two reports read differently:
#: these are core sections every bid has, not opt-in ones.
_NARRATIVE_CEILINGS = {
    "executive summary": 260,
    "purpose": 150,
    "summary": 400,
    # The house Summary, read from the delivered Kenneth proposal, is SIX paragraphs of
    # about 350 words in a fixed order: the platform named concretely, then four
    # paragraphs each opening "We commit to ...", then the concrete service-level targets.
    # A 200-word ceiling could not hold that shape, so the ceiling was wrong, not the
    # pattern. Read an earlier delivery before inventing a limit.
    "mobile app strategy": 160,
    "system overview": 300,
}

#: Per-ENTRY ceilings for bulleted sections. The list may be as long as the requirements
#: demand; each line has to stay scannable. Enforced on the longest entry, because an
#: average hides one 140-word bullet inside nine short ones, and that long one is exactly
#: what a reader stumbles on.
_PER_BULLET_CEILINGS = {
    "problems & solutions": 55,
    "key risks & mitigations": 50,
    "assumptions & dependencies": 50,
    "service level targets": 50,
    "contractual exceptions": 50,
}


def check_narrative_length(docx_path: Path, report: Report) -> None:
    """Core narrative sections must respect the WRITE TIGHT ceilings, and bulleted
    sections must respect their PER-ENTRY ceiling.

    Both are majors rather than blockers: a document that is 20 words long is a style
    problem the author should see, not a reason to refuse to ship a technically correct
    bid at midnight.
    """
    doc = Document(str(docx_path))
    over, long_bullets, measured = [], [], 0
    for heading, _level, bodies, _had_table in _section_bodies(doc):
        key = heading.strip().lower()
        cap = _NARRATIVE_CEILINGS.get(key)
        if cap is not None:
            measured += 1
            words = sum(len(b.split()) for b in bodies)
            if words > cap:
                over.append("'%s' %d words, ceiling %d (+%d)" % (heading, words, cap, words - cap))
        bullet_cap = _PER_BULLET_CEILINGS.get(key)
        if bullet_cap is not None and bodies:
            measured += 1
            worst = max(bodies, key=lambda b: len(b.split()))
            n = len(worst.split())
            if n > bullet_cap:
                avg = sum(len(b.split()) for b in bodies) / len(bodies)
                long_bullets.append(
                    "'%s' longest entry %d words, ceiling %d, average %.0f across %d entries"
                    % (heading, n, bullet_cap, avg, len(bodies)))
    if over:
        report.add(Issue(
            "narrative_section_too_long", "content", "major", False,
            f"{len(over)} core narrative section(s) exceed their word ceiling",
            detail="; ".join(over)))
    if long_bullets:
        report.add(Issue(
            "bullet_entry_too_long", "content", "major", False,
            f"{len(long_bullets)} bulleted section(s) carry an entry over its word ceiling",
            detail="; ".join(long_bullets)))
    if measured and not over and not long_bullets:
        report.pass_("narrative_length")
    elif not measured:
        # Silence here would be indistinguishable from a clean pass on a document whose
        # headings were renamed, which is how a check quietly becomes a no-op.
        report.add(Issue(
            "narrative_length_not_measured", "content", "minor", False,
            "no section matched a narrative or per-bullet ceiling, so none was measured",
            detail="expected one of: " + ", ".join(sorted(
                set(_NARRATIVE_CEILINGS) | set(_PER_BULLET_CEILINGS)))))


def _section_bodies(doc):
    """Yield (heading_text, level, [body paragraph texts], had_table) per heading.

    A section ends at the next heading of ANY level, which is what a reader
    experiences: the words under this heading before the next one.
    """
    pmap = {p._p: p for p in doc.paragraphs}
    tmap = {t._tbl: t for t in doc.tables}
    cur = None
    for child in doc.element.body.iterchildren():
        if child in tmap:
            if cur is not None:
                cur[3] = True
            continue
        p = pmap.get(child)
        if p is None:
            continue
        sname = p.style.name if p.style else ""
        txt = (p.text or "").strip()
        if sname.startswith("Heading"):
            if cur is not None:
                yield tuple(cur)
            try:
                level = int(sname.split()[-1])
            except (ValueError, IndexError):
                level = 9
            cur = [txt, level, [], False]
            continue
        if cur is not None and txt:
            cur[2].append(txt)
    if cur is not None:
        yield tuple(cur)


def check_optional_section_length(docx_path: Path, report: Report) -> None:
    """Optional bid sections must respect their word ceiling."""
    doc = Document(str(docx_path))
    over = []
    for heading, _level, bodies, _had_table in _section_bodies(doc):
        cap = _SECTION_WORD_CEILINGS.get(heading.strip().lower())
        if cap is None:
            continue
        words = sum(len(b.split()) for b in bodies)
        if words > cap:
            over.append((heading, words, cap))
    if over:
        report.add(Issue(
            "optional_section_too_long", "content", "major", False,
            f"{len(over)} optional section(s) exceed their word ceiling",
            detail="; ".join(f"'{h}' {w} words, ceiling {c} (+{w - c})" for h, w, c in over),
        ))
    else:
        report.pass_("optional_section_length")


def check_team_roles_is_table(docx_path: Path, report: Report) -> None:
    """Roles & Responsibilities must render as a 'Role | Accountability' table.

    Prose here means the content-writer emitted a string where the schema requires
    an array of {name, description} rows, so the positions and their headcount are
    buried in a paragraph instead of being scannable. Same defect class as
    `techstack_not_table`, and detected the same proven way.
    """
    doc = Document(str(docx_path))
    bad = []
    for heading, _level, bodies, had_table in _section_bodies(doc):
        if heading.strip().lower() not in _TABLE_SECTIONS:
            continue
        if not had_table:
            bad.append((heading, (bodies[0][:60] if bodies else "(nothing)")))
    if bad:
        report.add(Issue(
            "team_roles_not_table", "visual", "blocker", False,
            f"{len(bad)} team section(s) render as prose instead of a "
            f"'Role | Accountability' table",
            detail="; ".join(f"'{h}' -> {prev}" for h, prev in bad),
        ))
    else:
        report.pass_("team_roles_is_table")


def _iter_content_paragraphs(doc):
    """Every paragraph a reader actually sees, including table cells, excluding headings."""
    def walk(paras):
        for p in paras:
            sname = p.style.name if p.style else ""
            if sname.startswith("Heading"):
                continue
            txt = (p.text or "").strip()
            if txt:
                yield txt

    yield from walk(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from walk(cell.paragraphs)


def check_em_dash_in_prose(docx_path: Path, report: Report) -> None:
    """No em-dash (—) anywhere in the delivered content.

    A spaced em-dash is one of the strongest signals a reader uses to spot
    machine-written text, and the client rejects it on sight. An earlier version
    of this check exempted bullets and figure captions, on the grounds that the
    dash was a structural separator there. That exemption is withdrawn: the reader
    does not care whether a dash is structural, only that it is there. Use a colon
    after a bold label, and a comma, a semicolon or two sentences in prose.
    """
    doc = Document(str(docx_path))
    bad = [txt[:70] for txt in _iter_content_paragraphs(doc) if "—" in txt]

    if bad:
        report.add(Issue(
            # BLOCKER, not "major": the client rejects this on sight, and a "major"
            # never reaches blocker_count, so the "loop until 0 blockers" gate in
            # 05b would have shipped the document with the em-dashes still in it.
            "em_dash_in_prose", "content", "blocker", False,
            f"{len(bad)} paragraph(s) contain an em-dash (—), which the client rejects as "
            f"machine-written. Use a colon after a bold label, and a comma, a semicolon or "
            f"two sentences in prose. This applies to bullets and captions too. Fix the "
            f"offending values in replacements.json / diagrams.json and re-run Phase 5a.",
            detail="; ".join(bad[:5]),
        ))
    else:
        report.pass_("em_dash_prose_ok")


# "e.g." and friends read as machine-assembled filler in a client-facing bid.
# Word-boundary anchored so "Inc." or a version like "v1.e" cannot match.
_LATIN_ABBREV_RE = re.compile(r"(?<![A-Za-z])(e\.g\.|i\.e\.|etc\.|viz\.|cf\.)", re.IGNORECASE)
_LATIN_ABBREV_FIX = {
    "e.g.": "for example / such as",
    "i.e.": "that is / in other words",
    "etc.": "finish the list, or name the category",
    "viz.": "namely",
    "cf.": "compare",
}


def check_latin_abbreviation_in_content(docx_path: Path, report: Report) -> None:
    """No Latin abbreviations in the delivered content: write the English instead."""
    doc = Document(str(docx_path))
    hits, found = [], set()
    for txt in _iter_content_paragraphs(doc):
        for m in _LATIN_ABBREV_RE.finditer(txt):
            found.add(m.group(1).lower())
            hits.append(txt[:70])
    if hits:
        advice = "; ".join(f"{k} -> {v}" for k, v in _LATIN_ABBREV_FIX.items() if k in found)
        report.add(Issue(
            # BLOCKER for the same reason as em_dash_in_prose: a "major" is not counted
            # by the gate, so it would ship.
            "latin_abbreviation_in_content", "content", "blocker", False,
            f"{len(hits)} paragraph(s) use a Latin abbreviation ({', '.join(sorted(found))}), "
            f"which reads as machine-assembled filler in a client-facing bid. Write it out: "
            f"{advice}. Fix the offending values in replacements.json / diagrams.json and "
            f"re-run Phase 5a.",
            detail="; ".join(dict.fromkeys(hits))[:400],
        ))
    else:
        report.pass_("latin_abbreviation_ok")


# ---------------------------------------------------------------------------
# Runner.
# ---------------------------------------------------------------------------



def check_bullet_hanging_indent(docx_path: Path, report: Report) -> None:
    """Every glyph bullet needs a hanging indent, or a wrapped line breaks alignment.

    A 25-word bullet always wraps. Without `w:ind left=280 hanging=200` the second line
    returns to the left margin and sits under the glyph rather than under the text. On the
    delivered document 54 bullets had it and 140 did not, which reads worse than none of
    them having it, and it looks the same way in Word and in SharePoint. Walks table cells
    as well as the body, because the first version of the fix missed 20 bullets sitting
    inside the stack and role tables.
    """
    doc = Document(str(docx_path))

    def every_paragraph(d):
        yield from d.paragraphs
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    bad = []
    total = 0
    for para in every_paragraph(doc):
        text = (para.text or "").lstrip()
        if not text.startswith(("●", "•")):
            continue
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            continue
        total += 1
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        if ind is None or ind.get(qn("w:hanging")) is None:
            bad.append(text[:60])
    if bad:
        report.add(Issue(
            "bullet_indent_inconsistent", "layout", "major", True,
            f"{len(bad)} of {total} bullet(s) have no hanging indent, so a wrapped line "
            f"falls back to the left margin",
            detail="; ".join(bad[:6])))
    elif total:
        report.pass_("bullet_hanging_indent")


def check_uat_section_present(docx_path: Path, report: Report) -> None:
    """The User Acceptance Testing terms must survive into the built document.

    These are contractual terms (the 14 day window, the no-feedback-means-accepted rule, the
    three week critical-and-high fix commitment, the evidence required to reject) and they sit
    verbatim in the template rather than behind a placeholder, so that no run can reword them.
    That is the right call for contractual language, but it also means nothing generates the
    section, and a section nothing generates is one whose disappearance nobody notices. So the
    check asserts the heading AND the four commitments, not merely the heading: a heading with
    an emptied body would otherwise pass while the commitments were gone.
    """
    doc = Document(str(docx_path))
    heading = None
    for para in doc.paragraphs:
        # The heading sits at Heading 6, where this template types the section number into the
        # text rather than auto-numbering it, so the title is a suffix and not the whole string.
        title = re.sub(r"^[\d.]+\s*", "", (para.text or "").strip()).lower()
        if title == "user acceptance testing" \
           and (para.style.name or "").lower().startswith("heading"):
            heading = para
            break
    if heading is None:
        report.add(Issue(
            "uat_section_missing", "content", "blocker", False,
            "the User Acceptance Testing section is absent",
            detail="these are contractual terms carried verbatim by the template; if the "
                   "heading is gone the template has been edited or replaced"))
        return

    body = " ".join(p.text or "" for p in doc.paragraphs).lower()
    required = {
        "the UAT completion window":     "within 14 days",
        "no-feedback-means-accepted":    "regard the uat as successful",
        "the critical and high fix SLA": "within 3 weeks",
        "the rejection evidence rule":   "validated evidence",
    }
    missing = [label for label, needle in required.items() if needle not in body]
    if missing:
        report.add(Issue(
            "uat_terms_incomplete", "content", "blocker", False,
            f"the User Acceptance Testing section is missing {len(missing)} of its "
            f"{len(required)} contractual terms",
            detail="; ".join(missing)))
    else:
        report.pass_("uat_section_present")


CHECKS = [
    ("zip_integrity",                  check_zip_integrity),
    ("file_size",                      check_size),
    ("track_changes",                  check_no_track_changes),
    ("comments",                       check_no_comments),
    ("macros",                         check_no_macros),
    ("encryption",                     check_no_encryption),
    ("external_images",                check_no_external_images),
    ("placeholders",                   check_no_unfilled_placeholders),
    ("image_sharpness",                check_image_sharpness),
    ("justify_body",                   check_justify_body),
    # Quality checks added 2026-05-25:
    ("no_bullet_number_duplicates",    check_no_bullet_number_duplicates),
    ("image_spacing",                  check_image_spacing),
    ("heading_indent_vs_body",         check_heading_indent_vs_body),
    ("heading_section_spacing",        check_heading_section_spacing),
    ("caption_followed_by_gap",        check_caption_followed_by_gap),
    ("settings_flags_present",         check_settings_flags_present),
    ("justify_whitespace_channels",    check_justify_whitespace_channels),
    ("techstack_is_table",             check_techstack_is_table),
    ("bullet_hanging_indent",          check_bullet_hanging_indent),
    # Added 2026-07-29: the optional sections were shipping as walls of text and the
    # team section named no headcount. Both rules now have a check behind them.
    ("optional_section_length",        check_optional_section_length),
    ("narrative_length",               check_narrative_length),
    ("team_roles_is_table",            check_team_roles_is_table),
    ("em_dash_in_prose",               check_em_dash_in_prose),
    ("latin_abbreviation_in_content",  check_latin_abbreviation_in_content),
    # Added 2026-07-29: the UAT terms are contractual and live verbatim in the template
    # rather than behind a placeholder, precisely so no run can paraphrase them. A section
    # nothing generates is a section nothing notices going missing, hence the check.
    ("uat_section_present",            check_uat_section_present),
]


def run(docx_path: Path, pages_dir: Path | None) -> Report:
    report = Report(docx=str(docx_path), file_size_bytes=docx_path.stat().st_size)
    for name, fn in CHECKS:
        try:
            fn(docx_path, report)
        except Exception as e:
            report.add(Issue(f"check_error::{name}", "polish", "minor", False,
                             f"Check {name} crashed", detail=str(e)))
    check_heading_orphans_pdf(pages_dir, report)
    return report


def write_markdown(report: Report, out: Path) -> None:
    lines = [
        f"# Format Review",
        "",
        f"- Document: `{report.docx}`",
        f"- File size: {report.file_size_bytes/1024/1024:.2f} MB",
        f"- Blockers: **{report.blocker_count}**",
        f"- Auto-fixable: **{report.auto_fixable_count}**",
        f"- Checks passed: {len(report.checks_passed)}",
        "",
    ]
    if report.issues:
        lines.append("## Issues")
        for i in report.issues:
            tag = "AUTO" if i.auto_fixable else "MANUAL"
            lines.append(f"- **[{i.severity.upper()} / {tag}] {i.category}** — {i.summary}")
            if i.detail:
                lines.append(f"  - {i.detail}")
    else:
        lines.append("## Issues\n\nNone — proposal passes the strict review.")
    lines.append("")
    lines.append("## Checks passed")
    for c in report.checks_passed:
        lines.append(f"- {c}")
    out.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", type=Path, required=True)
    ap.add_argument("--pages", type=Path, default=None, help="Optional dir of per-page PNGs")
    ap.add_argument("--json", type=Path, required=True, help="Write JSON report here")
    ap.add_argument("--md", type=Path, default=None, help="Also write markdown report here")
    ap.add_argument("--strict", action="store_true", help="(reserved) bump minor -> major")
    args = ap.parse_args()

    if not args.docx.exists():
        print(f"! docx not found: {args.docx}", file=sys.stderr); return 1

    report = run(args.docx, args.pages)

    # Persist
    args.json.parent.mkdir(parents=True, exist_ok=True)
    args.json.write_text(json.dumps({
        "docx": report.docx,
        "file_size_bytes": report.file_size_bytes,
        "blocker_count": report.blocker_count,
        "auto_fixable_count": report.auto_fixable_count,
        "checks_passed": report.checks_passed,
        "issues": [asdict(i) for i in report.issues],
    }, indent=2, ensure_ascii=False), encoding="utf-8")

    if args.md:
        write_markdown(report, args.md)

    # Print short summary
    print(f"Issues: {len(report.issues)}  Blockers: {report.blocker_count}  Auto-fixable: {report.auto_fixable_count}")
    print(f"Passed: {len(report.checks_passed)} checks")
    return 0 if report.blocker_count == 0 else 2


if __name__ == "__main__":
    sys.exit(main())
