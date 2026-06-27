#!/usr/bin/env python3
"""Assemble a project's High-Level Technical Proposal .docx.

Inputs:
  --template       path to proposal_template.docx (the reusable template)
  --replacements   JSON dict: keys -> values to substitute in the template
  --diagrams       JSON list: each entry {slug, png, caption, target_heading}
  --out            output .docx path

Behaviour:
  1. Copy template to output path.
  2. Substitute every {{KEY}} token in body / header / footer / tables with
     the matching value from replacements.json, preserving run formatting
     where practical. (Cover-page client / project / version live in the
     template as {{KEY}} placeholders, so step 2 handles them too.)
  3. Inject the per-project H6 sub-headings under "System Overview" — one
     for each entry in diagrams.json, in order, placed after "2.1.1 System
     Context".
  4. For each diagram: find the heading whose text matches target_heading
     (case-insensitive contains match) and insert the PNG immediately after,
     followed by a Caption-styled paragraph.
  5. Drop empty optional sections (Mobile App Strategy if mobile_strategy is
     null; tech-stack Data/AI placeholders if those values are null).
  6. Justify body paragraphs and bullet text.
  7. Enable autoHyphenation + updateFields in word/settings.xml so the TOC
     refreshes the first time the user opens the file.
  8. Save.

The script is intentionally generic — no customer-, vendor-, or project-
specific strings live in this file.
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PLACEHOLDER_RE = re.compile(r"\{\{([A-Z_][A-Z0-9_]*)\}\}")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

# Anchor heading the dynamic H6 inject step (step 3) targets.
SUBHEADING_ANCHOR = "2.1.1 System Context"


def _render_inline_markdown_into(p_el, text: str) -> None:
    """Replace runs in p_el with new runs that interpret **bold** markdown.
    Existing runs are cleared first; the resulting paragraph keeps its
    paragraph properties (style, alignment, spacing) intact."""
    from docx.oxml import OxmlElement
    # Snapshot the first run's rPr (if any) so the new runs inherit base font.
    first_r = p_el.find(qn("w:r"))
    base_rPr = first_r.find(qn("w:rPr")) if first_r is not None else None
    # Drop all existing runs.
    for r in p_el.findall(qn("w:r")):
        p_el.remove(r)

    def add_run(segment: str, bold: bool) -> None:
        if segment == "":
            return
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        # Inherit base font properties (size, font family) from the cleared run.
        if base_rPr is not None:
            for child in base_rPr:
                tag = child.tag.rsplit("}", 1)[-1]
                if tag in ("b", "i"):
                    continue  # we set bold explicitly
                import copy as _copy
                rPr.append(_copy.deepcopy(child))
        if bold:
            rPr.append(OxmlElement("w:b"))
        if len(rPr):
            r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = segment
        r.append(t)
        p_el.append(r)

    # Split text on **...** markers.
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            add_run(text[pos:m.start()], bold=False)
        add_run(m.group(1), bold=True)
        pos = m.end()
    if pos < len(text):
        add_run(text[pos:], bold=False)


def _expand_paragraph_with_value(p_el, value: str) -> None:
    """Render `value` into the document at p_el's position, splitting on
    newlines so each line becomes its own paragraph, and interpreting
    `**bold**` markdown inside each line. Cloned paragraphs inherit the
    style + spacing of p_el so the visual rhythm stays consistent."""
    import copy as _copy
    lines = value.split("\n")
    if not lines:
        lines = [""]
    _render_inline_markdown_into(p_el, lines[0])
    cursor = p_el
    for line in lines[1:]:
        new_p = _copy.deepcopy(p_el)
        # Strip any sectPr from the clone so we don't duplicate section breaks.
        for sp in new_p.findall(".//" + qn("w:sectPr")):
            sp.getparent().remove(sp)
        _render_inline_markdown_into(new_p, line)
        cursor.addnext(new_p)
        cursor = new_p


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def style_id_to_name(doc) -> dict:
    return {s.style_id: s.name for s in doc.styles if getattr(s, "style_id", None)}


def is_heading(paragraph, name_map: dict) -> int | None:
    style = paragraph.style.name if paragraph.style else None
    if style and style.startswith("Heading"):
        try:
            return int(style.split()[-1])
        except ValueError:
            return None
    return None


def _stringify_value(val):
    """Render a replacement value (str / list / dict) into the multi-line
    markdown string that `_expand_paragraph_with_value` can lay out.

    - str           → as-is
    - list[dict] with problem/solution keys → 2 lines per entry
                                              ("● <problem>" / "**Solution:** <s>")
    - list[str]     → "● <item>" lines
    - list[other]   → "● <repr>" lines
    - dict          → "**key:** value" lines
    - None          → None (caller handles drop)
    """
    if val is None:
        return None
    if isinstance(val, str):
        return val
    if isinstance(val, list):
        if not val:
            return ""
        if all(isinstance(it, dict) and "problem" in it and "solution" in it for it in val):
            chunks = []
            for it in val:
                # Replace "**Bold Title** — body" with "**Bold Title**: body"
                # User feedback: em-dash after bold title in P&S bullets reads
                # AI-generated; colon is the preferred separator.
                prob = re.sub(r'(\*\*[^*]+\*\*)\s*—\s*', r'\1: ', it['problem'], count=1)
                sol  = re.sub(r'(\*\*[^*]+\*\*)\s*—\s*', r'\1: ', it['solution'], count=1)
                chunks.append(f"● {prob}")
                chunks.append(f"**Solution:** {sol}")
            return "\n".join(chunks)
        lines = []
        for it in val:
            if isinstance(it, str):
                lines.append(f"● {it}")
            elif isinstance(it, dict):
                lines.append("● " + "; ".join(f"**{k}:** {v}" for k, v in it.items()))
            else:
                lines.append(f"● {it}")
        return "\n".join(lines)
    if isinstance(val, dict):
        return "\n".join(f"**{k}:** {v}" for k, v in val.items())
    return str(val)


def replace_in_paragraph(paragraph, mapping: dict[str, str]) -> int:
    """Replace {{KEY}} tokens in a paragraph's runs. Returns count replaced.

    We work on the joined text first, then write the result back into the
    first run and clear the others. This loses per-token formatting inside
    the affected runs but preserves the paragraph style. For the proposal
    template this is acceptable — placeholders sit inside Normal text.
    """
    joined = "".join(r.text or "" for r in paragraph.runs)
    if not joined:
        return 0

    count = 0

    def sub(m):
        nonlocal count
        key = m.group(1).lower()
        if key in mapping:
            val = mapping[key]
            if val is None:
                # Leave the placeholder so drop_section_if_empty can remove the
                # surrounding heading + body cleanly. str(None) would silently
                # litter the document with the word "None".
                return m.group(0)
            count += 1
            return _stringify_value(val)
        return m.group(0)

    new = PLACEHOLDER_RE.sub(sub, joined)

    if new == joined:
        return 0

    runs = paragraph.runs
    if not runs:
        return 0
    # If the replaced text carries multi-paragraph content (\n) or bold
    # markdown (**...**), render it properly into the document instead of
    # collapsing into one runs[0].text. Otherwise fall back to the simple
    # path — keeps run formatting for plain substitutions.
    if "\n" in new or "**" in new:
        _expand_paragraph_with_value(paragraph._p, new)
        return count
    runs[0].text = new
    for r in runs[1:]:
        r.text = ""
    return count


def replace_in_table(table, mapping: dict[str, str]) -> int:
    total = 0
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                total += replace_in_paragraph(p, mapping)
            for nested in cell.tables:
                total += replace_in_table(nested, mapping)
    return total


def replace_in_section_parts(section, mapping: dict[str, str]) -> int:
    """Replace in header / footer (default + first-page + even)."""
    total = 0
    for part in (section.header, section.footer, section.first_page_header,
                 section.first_page_footer, section.even_page_header, section.even_page_footer):
        try:
            for p in part.paragraphs:
                total += replace_in_paragraph(p, mapping)
            for t in part.tables:
                total += replace_in_table(t, mapping)
        except Exception:
            # Section may not have all part types; ignore.
            pass
    return total


def clone_paragraph_after(anchor_p, new_text: str):
    """Clone the anchor paragraph (style + format preserved), set its text, and
    insert it immediately after the anchor. Returns the new paragraph element."""
    import copy as _copy
    new_el = _copy.deepcopy(anchor_p._p)
    runs = new_el.findall(qn("w:r"))
    set_once = False
    for r in runs:
        for t in r.findall(qn("w:t")):
            if not set_once:
                t.text = new_text
                set_once = True
            else:
                t.text = ""
    if not set_once and runs:
        from docx.oxml import OxmlElement
        t = OxmlElement("w:t")
        t.text = new_text
        runs[0].append(t)
    anchor_p._p.addnext(new_el)
    return new_el


def inject_subheadings_after_anchor(doc, anchor_text: str, subheadings: list[str]) -> int:
    """Insert one H6 paragraph per item in `subheadings` after the anchor heading.

    Auto-renumbers each injected sub-heading sequentially starting from the
    anchor's number + 1, so there are no gaps even if upstream content
    skipped or reordered entries. Stripped from each input string:
    any existing 'N.N.N' prefix.
    """
    anchor = find_heading_paragraph(doc, anchor_text)
    if anchor is None or not subheadings:
        return 0

    # Parse anchor's numbering ("2.1.1 System Context" -> prefix "2.1.", next 2)
    m = re.match(r"\s*(\d+(?:\.\d+)*)\.(\d+)\s+", anchor_text)
    if m:
        prefix = m.group(1) + "."
        start = int(m.group(2)) + 1
    else:
        prefix = ""
        start = 1

    renumbered = []
    for i, text in enumerate(subheadings):
        clean = re.sub(r"^\s*\d+(\.\d+)*\s*[.:)-]?\s*", "", text).strip()
        renumbered.append(f"{prefix}{start + i} {clean}" if prefix else clean)

    # Insert in reverse so document order matches input order.
    for text in reversed(renumbered):
        clone_paragraph_after(anchor, text)
    return len(renumbered)


def find_heading_paragraph(doc, target: str):
    """Find the first heading paragraph whose text contains `target` (case-insensitive)."""
    target_norm = target.lower().strip()
    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            if target_norm in p.text.lower():
                return p
    return None


def insert_after_paragraph(paragraph, new_element) -> None:
    """Insert raw XML element immediately after `paragraph` in the body."""
    paragraph._p.addnext(new_element)


def _build_seq_caption_table(doc, title: str):
    """Clone an existing 1x1 caption table from the document and put
    'Figure <SEQ Figure>: <title>' inside its single cell. Returns the new
    table element ready to be inserted into the body.

    Cloning preserves the template's bordered/shaded cell style so every
    figure in the document — verbatim Section 3 plus Phase-4-inserted
    Section 2 diagrams — looks identical."""
    import copy as _copy
    from docx.oxml import OxmlElement

    # Find a sample 1x1 caption table in the doc to clone.
    sample = None
    for t in doc.tables:
        if len(t.rows) == 1 and len(t.columns) == 1:
            txt = t.rows[0].cells[0].text.strip()
            if txt:
                sample = t
                break
    if sample is None:
        return None

    new_tbl = _copy.deepcopy(sample._tbl)
    # Wipe existing text in the clone.
    for tn in new_tbl.iter(qn("w:t")):
        tn.text = ""
    # Find the target paragraph inside the single cell.
    cell_p = next(iter(new_tbl.iter(qn("w:p"))), None)
    if cell_p is None:
        return new_tbl
    # Take the first existing run's rPr as the formatting template for SEQ runs.
    first_r = cell_p.find(qn("w:r"))
    rpr_tmpl = first_r.find(qn("w:rPr")) if first_r is not None else None

    def text_run(text: str):
        r = OxmlElement("w:r")
        if rpr_tmpl is not None:
            r.append(_copy.deepcopy(rpr_tmpl))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        return r

    def fld_run(ftype: str):
        r = OxmlElement("w:r")
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), ftype)
        if ftype == "begin":
            fc.set(qn("w:dirty"), "true")
        r.append(fc)
        return r

    def instr_run(text: str):
        r = OxmlElement("w:r")
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = text
        r.append(it)
        return r

    # Remove any leftover runs so the cell paragraph starts empty.
    for r in cell_p.findall(qn("w:r")):
        cell_p.remove(r)
    cell_p.append(text_run("Figure "))
    cell_p.append(fld_run("begin"))
    cell_p.append(instr_run(r" SEQ Figure \* ARABIC "))
    cell_p.append(fld_run("separate"))
    cell_p.append(text_run("1"))
    cell_p.append(fld_run("end"))
    cell_p.append(text_run(f": {title}"))
    return new_tbl


# Patterns the content-writer agent emits at the start of bullet items
# that we MUST strip so the rendered bullet doesn't read "●  1. text" or
# "●  Step 1 — text" (the user-flagged duplicate-numbering bug).
#
# Critically, we must NOT destroy the bold component label. Inputs like
#   "1. **Step 1 — Bank API POST** — body"
# should become
#   "**Bank API POST** — body"
# — i.e. strip ONLY the manual numbering, including the "Step 1 — " prefix
# inside the bold span. The bold-wrapped title text must survive.

# Plain manual numbering at the start of the string (no surrounding emphasis).
_PLAIN_NUMBER_RE = re.compile(
    r"^\s*(?:"
    r"\(?\d{1,3}\)?\s*[.)\-:]\s+"              # "1. "  "(1) "  "1) "  "1 - "
    r"|step\s+\d{1,3}[a-z]?\s*[.)\-:—]\s+"     # "Step 1 — "  "Step 2a — "
    r")",
    re.IGNORECASE,
)

# Manual numbering INSIDE a leading bold/italic span. We match the opening
# emphasis marker + the numbering, then re-emit the opening marker so the
# title text immediately following stays bold. Example:
#   "**Step 1 — Bank API**"
# becomes (after re.sub):
#   "**Bank API**"
_INSIDE_EMPH_NUMBER_RE = re.compile(
    r"^(\s*)(\*\*|__|\*|_)\s*"
    r"(?:"
    r"\(?\d{1,3}\)?\s*[.)\-:]\s+"
    r"|step\s+\d{1,3}[a-z]?\s*[.)\-:—]\s+"
    r")",
    re.IGNORECASE,
)


def _strip_manual_numbering(text: str) -> str:
    """Remove ALL leading manual numbering from a bullet item while
    preserving the bold component label.

    Two patterns per pass:
      1. "1. " / "Step 1 — " at the very start → strip outright
      2. "**Step 1 — title**" / "**1. title**" → strip ONLY the numbering
         from inside the bold span, leaving "**title**" intact

    Loops because the content-writer agent stacks (e.g. "1. **Step 1 — title**")
    — one pass strips the leading "1. ", the next strips the "Step 1 — "
    from inside the bold span.
    """
    if not text:
        return text
    stripped = text.lstrip()
    while stripped.startswith("●"):
        stripped = stripped[1:].lstrip()
    for _ in range(4):
        # Strip "**Step 1 — " → "**" first so the title stays bold-wrapped.
        m = _INSIDE_EMPH_NUMBER_RE.match(stripped)
        if m:
            stripped = m.group(1) + m.group(2) + stripped[m.end():]
            stripped = stripped.lstrip()
            continue
        m = _PLAIN_NUMBER_RE.match(stripped)
        if m:
            stripped = stripped[m.end():].lstrip()
            continue
        break
    return stripped.strip()


def _build_bullet_paragraph(text: str, *, is_first: bool = False):
    """Construct a justified body paragraph for a diagram explanation bullet.
    Strips any manual numbering the agent prefixed (avoids "● 1. text"
    duplication), prepends "●  ", then renders `**bold**` inline markdown.

    is_first=True: this is the bullet immediately after a Figure caption
    table. The table cannot carry `space_after`, so the first bullet must
    declare a larger `space_before` to create the breathing gap between
    caption and first bullet that the user has flagged as missing.
    """
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "240" if is_first else "0")  # 12pt before first
    spacing.set(qn("w:after"), "120")
    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    pPr.append(jc)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "280")     # tighter left indent (was 360)
    ind.set(qn("w:hanging"), "200")  # bullet sits at 80 twips, text wraps at 280
    pPr.append(ind)
    p.append(pPr)
    # Single space between bullet glyph and text (was 2) — user flagged
    # the gap as "bullet vs chữ xa nhau".
    body_text = "● " + _strip_manual_numbering(text)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = ""
    r.append(t)
    p.append(r)
    _render_inline_markdown_into(p, body_text)
    return p


def _build_intro_paragraph(text: str):
    """Justified body paragraph (intro sentence between H6 and image).
    Renders `**bold**` markdown inside the text."""
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "120")
    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    pPr.append(jc)
    p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = ""
    r.append(t)
    p.append(r)
    _render_inline_markdown_into(p, text)
    return p


def insert_image_after_heading(doc, heading_text: str, png_path: Path,
                                caption: str, width_inches: float = 6.5,
                                intro_paragraph: str | None = None,
                                explanation_bullets: list[str] | None = None) -> bool:
    """Insert (in document order):

        heading
        → intro_paragraph        (optional context sentence)
        → image                  (centered)
        → SEQ-numbered Figure caption table
        → explanation bullet 1
        → explanation bullet 2
        → …

    This matches the standard STS proposal convention: every diagram has
    an intro sentence above it and 7–17 explanation bullets below it.
    """
    heading = find_heading_paragraph(doc, heading_text)
    if heading is None:
        return False

    body = doc.element.body
    placeholder_p = doc.add_paragraph()
    placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add breathing room above + below the figure. Without this the image
    # touches the intro sentence above and the bullets below — a complaint
    # the user has flagged as a quality issue. 12pt ≈ one body line.
    placeholder_p.paragraph_format.space_before = Pt(12)
    placeholder_p.paragraph_format.space_after = Pt(6)  # caption sits below; small gap to it
    placeholder_run = placeholder_p.add_run()
    placeholder_run.add_picture(str(png_path), width=Inches(width_inches))

    title = re.sub(r"^\s*Figure\s+\d+\s*[:.]?\s*", "", caption).strip() or caption
    caption_el = _build_seq_caption_table(doc, title)

    body.remove(placeholder_p._p)
    # Insert in REVERSE so document order ends up: heading → intro → image
    # → caption → bullets (in list order). The LAST item we insert (which
    # appears FIRST in document order, immediately after the caption table)
    # carries is_first=True so it owns the space-before gap that the caption
    # table can't provide itself.
    bullets = list(explanation_bullets or [])
    non_empty = [b for b in bullets if b and str(b).strip()]
    for idx, bullet_text in enumerate(reversed(non_empty)):
        is_first_in_list = (idx == len(non_empty) - 1)
        insert_after_paragraph(heading,
                               _build_bullet_paragraph(str(bullet_text),
                                                       is_first=is_first_in_list))
    if caption_el is not None:
        insert_after_paragraph(heading, caption_el)
    insert_after_paragraph(heading, placeholder_p._p)
    if intro_paragraph and intro_paragraph.strip():
        insert_after_paragraph(heading, _build_intro_paragraph(intro_paragraph.strip()))
    return True


def drop_section_if_empty(doc, heading_text: str, mapping: dict, content_key: str | None = None) -> bool:
    """If the value for the section's content key is null/empty, remove the heading
    and any following paragraphs up to the next heading at same-or-higher level.
    Returns True if section was dropped.

    `content_key`, when given, is the exact mapping key to test (use this when the
    heading TEXT no longer equals the key, e.g. a "Data" heading whose content key
    is "techstack_data"). When omitted, the key is derived from the heading text:
    lowercased, spaces/dashes/braces stripped ("Mobile App Strategy" ->
    "mobile_app_strategy").
    """
    heading = find_heading_paragraph(doc, heading_text)
    if heading is None:
        return False

    if content_key is not None:
        key = content_key
    else:
        key = heading_text.lower()
        key = re.sub(r"[{}]", "", key)
        key = re.sub(r"[\s\-]+", "_", key).strip("_")
    # Empty == None / "" / "null" / [] (an empty techstack array is also "no section").
    if mapping.get(key) not in (None, "", "null", []):
        return False

    own_level = int(heading.style.name.split()[-1])
    to_remove = [heading._p]
    cursor = heading._p
    while True:
        nxt = cursor.getnext()
        if nxt is None:
            break
        tag = nxt.tag.rsplit("}", 1)[-1]
        if tag == "p":
            # Check if this is another heading at same-or-higher level
            pPr = nxt.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    val = pStyle.get(qn("w:val"), "")
                    # Need style name from styles part. Approximate: heading style ids
                    # in this template are numeric — bail out conservatively at any heading.
                    # (Caller responsibility: design templates so optional sections are leaf headings.)
                    name = doc.styles.element.findall(qn("w:style"))
                    # Quick check: if val matches a heading style by name, stop.
                    style_name = next((s.get(qn("w:styleId")) for s in name if s.get(qn("w:styleId")) == val), None)
                    # Fallback: stop at the next paragraph that python-docx reports as a heading.
                    for p in doc.paragraphs:
                        if p._p is nxt and p.style and p.style.name.startswith("Heading"):
                            lvl = int(p.style.name.split()[-1])
                            if lvl <= own_level:
                                # Stop before this paragraph
                                for el in to_remove:
                                    el.getparent().remove(el)
                                return True
                            break
            to_remove.append(nxt)
        elif tag == "tbl":
            to_remove.append(nxt)
        else:
            # sectPr etc — stop
            break
        cursor = nxt

    for el in to_remove:
        el.getparent().remove(el)
    return True


# Visual hierarchy: bigger gap before a heading than after it, so the
# heading reads as introducing the block that follows rather than floating
# loose between two equal whitespace gulfs. Values are in points and are
# floors — if the template already has more, we leave it alone.
# Floors MUST match the strict reviewer's heading_section_spacing_tight floors
# (H1-H3 >= 24pt, H4-H5 >= 18pt, H6 >= 12pt space-before) so that re-running
# build_docx after auto_fix does NOT clamp headings back below the reviewer
# floor and re-introduce the defect. Kept in sync with format_reviewer.py.
_HEADING_MIN_SPACE_BEFORE = {1: 24, 2: 24, 3: 24, 4: 18, 5: 18, 6: 12}
_HEADING_MIN_SPACE_AFTER  = {1:  6, 2:  6, 3:  6, 4:  5, 5:  5, 6:  4}
# Hard CEILING — keeps section transitions visible without going loose; must
# sit ABOVE the reviewer floors above so the ceiling never clamps below them.
_HEADING_MAX_SPACE_BEFORE = {1: 36, 2: 36, 3: 36, 4: 28, 5: 28, 6: 20}
_HEADING_MAX_SPACE_AFTER  = {1: 14, 2: 14, 3: 14, 4: 12, 5: 12, 6: 10}


def boost_heading_spacing(doc) -> int:
    """Raise heading space-before / space-after to the floors above. Without
    this Heading 5/6 sit too close to their preceding paragraph and the
    document reads as a wall of text — a quality issue the user has
    flagged. Floors only: any template value above the floor is preserved.
    """
    touched = 0
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if not sname.startswith("Heading"):
            continue
        try:
            lvl = int(sname.split()[-1])
        except ValueError:
            continue
        sb_floor = _HEADING_MIN_SPACE_BEFORE.get(lvl)
        sa_floor = _HEADING_MIN_SPACE_AFTER.get(lvl)
        sb_ceil = _HEADING_MAX_SPACE_BEFORE.get(lvl)
        sa_ceil = _HEADING_MAX_SPACE_AFTER.get(lvl)
        if sb_floor is None:
            continue
        pf = p.paragraph_format
        cur_sb = pf.space_before.pt if pf.space_before else 0
        cur_sa = pf.space_after.pt if pf.space_after else 0
        # Floor: too tight reads as wall of text
        if cur_sb < sb_floor:
            pf.space_before = Pt(sb_floor)
            touched += 1
        if cur_sa < sa_floor:
            pf.space_after = Pt(sa_floor)
            touched += 1
        # Ceiling: too loose reads as orphan headings with hollow gaps
        if sb_ceil and cur_sb > sb_ceil:
            pf.space_before = Pt(sb_ceil)
            touched += 1
        if sa_ceil and cur_sa > sa_ceil:
            pf.space_after = Pt(sa_ceil)
            touched += 1
    return touched


_HEADING_TEXT_RE = re.compile(r"^\s*\d+(\.\d+){0,4}\.?\s+\S")


def _looks_like_heading_text(text: str) -> bool:
    """Heuristic: a paragraph whose visible text starts with `1.`, `1.1`,
    `2.1.3.`, etc. followed by a real word is structurally a heading even
    if the template forgot to apply the Heading 1/2/3 style. The proposal
    template carries these as plain paragraphs with hard-coded left_indent
    (480 / 960 / 1200 twentieths), so `pStyle.startswith("Heading")` misses
    them — which is the root cause of the "2.1 lòi ra ngoài 2 nằm trong"
    defect."""
    return bool(_HEADING_TEXT_RE.match(text or ""))


def normalize_heading_indents(doc) -> int:
    """Pull every heading paragraph back to flush-left so that:
      - the auto-number / typed number anchor sits at column 0 (the page
        margin), and
      - no heading sticks left of the body (the "2.1 lòi ra ngoài" defect),
        and
      - all headings line up consistently rather than stair-stepping
        deeper for each level.

    We catch headings via TWO signals (either is sufficient):
      1. Paragraph style starts with "Heading" (the python-docx-native way), OR
      2. The visible text matches `^\\d+(\\.\\d+)*\\.?\\s+\\S` — i.e. starts
         with `1.`, `1.1`, `2.1.3.`, etc. This catches the template's
         numbered-paragraph-without-Heading-style case.

    For every match we set BOTH left_indent and first_line_indent to 0 —
    leaving a hanging gutter would push the number off the left margin.
    """
    from docx.oxml import OxmlElement
    touched = 0
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        text = p.text or ""
        is_heading_style = sname.startswith("Heading")
        is_heading_text = _looks_like_heading_text(text)
        if not (is_heading_style or is_heading_text):
            continue
        pf = p.paragraph_format
        # ALWAYS force left_indent and first_line_indent to 0 on a heading.
        # Setting to None (the python-docx "inherit from style") is NOT
        # enough — the style or numbering definition still pushes the
        # number off-margin. We need a HARD zero at the paragraph level.
        pf.left_indent = Pt(0)
        pf.first_line_indent = Pt(0)
        touched += 1
        # Also write a direct XML w:ind override so the numbering
        # definition's indent attributes do not win over the paragraph's
        # explicit value. We FORCE-INSERT an <w:ind/> element with all
        # four ind attrs zeroed; this overrides the numId-derived indent.
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._p.insert(0, pPr)
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            # Insert ind in correct schema position (after numPr if present)
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                numPr.addnext(ind)
            else:
                pPr.append(ind)
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:start"), "0")
        ind.set(qn("w:firstLine"), "0")
        ind.set(qn("w:hanging"), "0")
    return touched


def ensure_caption_followed_by_gap(doc, min_pt: int = 12) -> int:
    """For every 1x1 caption table in the document, make sure the paragraph
    immediately after has space-before >= min_pt. A caption table cannot
    carry trailing whitespace itself, so the next paragraph must own the
    gap — otherwise body text sits flush against the caption, which the
    user has flagged as a layout defect ("text và title image vẫn còn quá
    sát nhau").
    """
    from docx.oxml import OxmlElement
    touched = 0
    body = doc.element.body
    for tbl in list(doc.tables):
        rows = tbl.rows
        if len(rows) != 1 or len(rows[0].cells) != 1:
            continue
        cell_text = "".join(p.text for p in rows[0].cells[0].paragraphs).strip()
        if not cell_text:
            continue  # not a caption (empty)
        # Skip if this isn't really a caption (heuristic: caption text contains
        # 'Figure' or starts with a number reference). Otherwise we'd push
        # body table-cell content around.
        if "Figure" not in cell_text and "figure" not in cell_text and "Bảng" not in cell_text:
            continue
        nxt = tbl._element.getnext()
        if nxt is None or nxt.tag.rsplit("}", 1)[-1] != "p":
            continue
        # Read current space-before
        pPr = nxt.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            nxt.insert(0, pPr)
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        cur_before = spacing.get(qn("w:before"))
        cur_pt = int(cur_before) / 20 if cur_before else 0
        if cur_pt < min_pt:
            spacing.set(qn("w:before"), str(min_pt * 20))  # twentieths of a point
            touched += 1
    return touched


def justify_body(doc) -> int:
    """Set every Normal / List* paragraph to JUSTIFY. Return count touched.

    SKIPS paragraphs that contain an inline image (`<w:drawing>`) — those
    are already centered by `center_image_paragraphs` and overriding to
    JUSTIFY would un-center the image (the previous bug 3.5.3 surfaced).
    """
    touched = 0
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if sname.startswith("Heading") or sname in ("Title", "Subtitle", "Caption"):
            continue
        # Skip image-bearing paragraphs — keep their CENTER alignment.
        if p._p.find(".//" + qn("w:drawing")) is not None:
            continue
        if p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            touched += 1
    return touched


def replace_remaining_placeholders_xml(out_path: Path, mapping: dict) -> int:
    """Final XML-level pass to catch {{KEY}} placeholders that python-docx
    iteration misses — notably TOC field cached text, comments XML, and any
    other parts that aren't standard body paragraphs/tables/headers.

    Only substitutes when the mapping value is not None; None values mean
    the surrounding section was supposed to be dropped, so leaving the
    placeholder is intentional and the format reviewer will flag it for the
    user to fix the upstream replacements file.
    """
    tmp = out_path.with_suffix(".tmp.docx")
    pattern = re.compile(r"\{\{([A-Z_][A-Z0-9_]*)\}\}")
    total = 0
    with zipfile.ZipFile(out_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".xml"):
                txt = data.decode("utf-8", errors="replace")
                def repl(m):
                    nonlocal total
                    key = m.group(1).lower()
                    if key in mapping and mapping[key] is not None:
                        total += 1
                        rendered = _stringify_value(mapping[key])
                        return rendered if rendered is not None else m.group(0)
                    return m.group(0)
                new_txt = pattern.sub(repl, txt)
                if new_txt != txt:
                    data = new_txt.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(out_path)
    return total


def patch_numbering_xml(out_path: Path) -> int:
    """Force every heading-numbering level to left-align (`lvlJc="left"`).

    The proposal_template's numbering definition uses `lvlJc="right"` so the
    auto-numbers (1., 1.1., 1.1.1.) are right-aligned within their hanging
    gutter. Because "1.1." is wider than "1.", a right-aligned 1.1. starts
    LEFT of where a right-aligned 1. starts — producing the "mục con nằm
    ngoài mục cha" defect the user has flagged (page 4 of the rendered docx).
    Left-aligning the numbers anchors them at the same left column.
    """
    tmp = out_path.with_suffix(".tmp.docx")
    touched = 0
    with zipfile.ZipFile(out_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/numbering.xml":
                txt = data.decode("utf-8")
                # Replace every "lvlJc" right-align inside any heading <w:lvl>
                new = re.sub(r'<w:lvlJc w:val="right"/>',
                             '<w:lvlJc w:val="left"/>', txt)
                touched += txt.count('<w:lvlJc w:val="right"/>')
                if new != txt:
                    data = new.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(out_path)
    return touched


def renumber_figure_sequence_cache(out_path: Path) -> int:
    """Walk word/document.xml and renumber every SEQ Figure field's cached
    text run with an incrementing integer 1, 2, 3, ... in document order.

    Why: viewers that DON'T evaluate Word field codes (LibreOffice PDF
    export, PyMuPDF, SharePoint preview) display the CACHED text between
    `<w:fldChar w:fldCharType="separate"/>` and `<w:fldChar w:fldCharType="end"/>`
    instead of evaluating `SEQ Figure`. If every caption ships with the same
    cached "1", every preview shows "Figure 1: ..." regardless of position.
    Renumbering the cache makes captions correct in EVERY viewer, before
    the user even presses F9.
    """
    tmp = out_path.with_suffix(".tmp.docx")
    seq_count = 0
    with zipfile.ZipFile(out_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                txt = data.decode("utf-8")
                # Pattern: <w:instrText ...>...SEQ Figure...</w:instrText>
                # followed by <w:fldChar w:fldCharType="separate"/>
                # followed by one or more runs until <w:fldChar w:fldCharType="end"/>
                # We rewrite the FIRST <w:t>...</w:t> inside that span.
                counter = {"n": 0}
                def renum_match(m):
                    counter["n"] += 1
                    seq_block = m.group(0)
                    # Find the cached text run between separate and end
                    new_block = re.sub(
                        r'(<w:fldChar w:fldCharType="separate"/></w:r>.*?<w:t[^>]*>)([^<]*)(</w:t>)',
                        lambda mm: mm.group(1) + str(counter["n"]) + mm.group(3),
                        seq_block,
                        count=1,
                        flags=re.DOTALL,
                    )
                    return new_block
                pattern = re.compile(
                    r'<w:instrText[^>]*>\s*SEQ\s+Figure[^<]*</w:instrText>.*?<w:fldChar w:fldCharType="end"/>',
                    re.DOTALL,
                )
                new_txt = pattern.sub(renum_match, txt)
                seq_count = counter["n"]
                if new_txt != txt:
                    data = new_txt.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(out_path)
    return seq_count


def patch_settings_xml(out_path: Path) -> None:
    """Edit word/settings.xml to add the document-level flags that keep our
    output looking clean. We add (only if missing):

      - autoHyphenation=true       — lets Word break long words across lines
                                     instead of leaving giant trailing gaps
                                     in justified paragraphs.
      - consecutiveHyphenLimit=2   — prevents more than 2 consecutive hyphens
                                     (Word default = unlimited, which can produce
                                     ugly hyphen ladders).
      - hyphenationZone=288        — 0.2" (in twentieths-of-a-point: 288);
                                     narrower zone → more hyphenation, fewer gaps.
      - doNotExpandShiftReturn=true — soft line breaks (Shift+Enter) do NOT
                                     trigger full justify expansion on the line
                                     that ends with one. This is the single
                                     most effective fix for the "huge white
                                     channels in middle of paragraph" symptom
                                     the user has reported.
      - characterSpacingControl=compressPunctuation — lets Word squeeze inter-
                                     character whitespace around punctuation
                                     instead of widening word gaps.
      - updateFields=true          — refreshes the TOC the first time the
                                     document is opened.
    """
    tmp = out_path.with_suffix(".tmp.docx")
    desired = {
        "autoHyphenation":            '<w:autoHyphenation w:val="true"/>',
        "consecutiveHyphenLimit":     '<w:consecutiveHyphenLimit w:val="2"/>',
        "hyphenationZone":            '<w:hyphenationZone w:val="288"/>',
        "doNotExpandShiftReturn":     '<w:doNotExpandShiftReturn/>',
        "characterSpacingControl":    '<w:characterSpacingControl w:val="compressPunctuation"/>',
        "updateFields":               '<w:updateFields w:val="true"/>',
    }
    with zipfile.ZipFile(out_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                txt = data.decode("utf-8")
                inserts = [xml for tag, xml in desired.items() if tag not in txt]
                if inserts:
                    txt = txt.replace(
                        "</w:settings>",
                        "".join(inserts) + "</w:settings>",
                    )
                data = txt.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(out_path)


SKILL_ROOT = Path(__file__).resolve().parent.parent
ICONS_ROOT = SKILL_ROOT / "assets" / "icons"


def _resolve_logo_path(logo_ref: str | None) -> Path | None:
    """Resolve a logo reference like 'azure/aks' or 'aws/lambda' to a real PNG."""
    if not logo_ref:
        return None
    try:
        pack, name = logo_ref.split("/", 1)
    except ValueError:
        return None
    candidates = [
        ICONS_ROOT / pack / "png" / f"{name}.png",
        ICONS_ROOT / pack / "png" / f"{name}@300.png",
    ]
    for c in candidates:
        if c.exists():
            return c
    return None


def _set_cell_fill(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_table_borders(table, color: str = "B7B7B7", size: int = 8) -> None:
    from docx.oxml import OxmlElement
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        return
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:color"), color)
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "0")
        borders.append(b)
    # Remove existing borders if any
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _build_techstack_table(doc, items: list[dict], header_fill: str = "595959"):
    """Build a 2-col 'Technology | Advantages' table styled like Tay Ho.

    Each item: {"name": str, "logo": "pack/name" | None, "description": str}.
    Logo (if resolvable) is embedded above the bold name in column 1.
    """
    from docx.shared import Inches as _Inches, RGBColor
    from docx.enum.table import WD_ALIGN_VERTICAL
    n_rows = 1 + len(items)
    tbl = doc.add_table(rows=n_rows, cols=2)
    tbl.autofit = False
    # Column widths — Technology col carries only logo + short name, keep narrow.
    # We MUST set both tblGrid (authoritative for Word layout) AND per-row cell
    # widths — python-docx only does the latter, leaving tblGrid stale.
    # Body width of this template's main section is 7.27" (10466 twips,
    # margins 0.5" + 0.5"). Table fills body margin-to-margin.
    col0_w = _Inches(1.75)
    col1_w = _Inches(5.52)
    # 1.75" = 2520 twips, 5.52" = 7946 twips, total 10466 = body width
    from docx.oxml import OxmlElement
    tblGrid = tbl._element.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for child in list(tblGrid):
            tblGrid.remove(child)
        for w in (2520, 7946):
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(w))
            tblGrid.append(gc)
    # Force table width = full page body (9360 twips) and fixed layout
    # so column widths take effect immediately. Also zero the table indent
    # so the table starts at the left page margin, ends at right margin.
    tblPr = tbl._element.find(qn("w:tblPr"))
    if tblPr is not None:
        # Width
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "10466")
        tblW.set(qn("w:type"), "dxa")
        # Indent (zero — flush with body left margin)
        tblInd = tblPr.find(qn("w:tblInd"))
        if tblInd is None:
            tblInd = OxmlElement("w:tblInd")
            tblPr.append(tblInd)
        tblInd.set(qn("w:w"), "0")
        tblInd.set(qn("w:type"), "dxa")
        # Fixed layout — column widths honored exactly, not autofit
        tblLayout = tblPr.find(qn("w:tblLayout"))
        if tblLayout is None:
            tblLayout = OxmlElement("w:tblLayout")
            tblPr.append(tblLayout)
        tblLayout.set(qn("w:type"), "fixed")
    for row in tbl.rows:
        row.cells[0].width = col0_w
        row.cells[1].width = col1_w
        # Vertical-center every cell so logo + name sit centered (user feedback "logo chữ nằm giữa")
        for c in row.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Header
    hdr = tbl.rows[0]
    for col_idx, label in enumerate(("Technology", "Advantages")):
        cell = hdr.cells[col_idx]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        _set_cell_fill(cell, header_fill)

    # Data rows
    for i, item in enumerate(items):
        row = tbl.rows[i + 1]
        # Col 1: optional logo + bold name, both centered horizontally
        name_cell = row.cells[0]
        name_cell.text = ""  # clear default empty paragraph
        # Logo paragraph (centered)
        logo_path = _resolve_logo_path(item.get("logo"))
        if logo_path:
            logo_para = name_cell.paragraphs[0]
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run()
            try:
                logo_run.add_picture(str(logo_path), height=_Inches(0.35))
            except Exception:
                pass
            # Name paragraph below logo (centered)
            name_para = name_cell.add_paragraph()
        else:
            name_para = name_cell.paragraphs[0]
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(item.get("name", ""))
        name_run.bold = True
        name_run.font.size = Pt(11)

        # Col 2: description
        desc_cell = row.cells[1]
        desc_cell.text = ""
        desc_para = desc_cell.paragraphs[0]
        _render_inline_markdown_into(desc_para._p, item.get("description", ""))
        desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _set_table_borders(tbl, color="B7B7B7", size=8)
    return tbl


def render_techstack_tables(doc, replacements: dict) -> set:
    """For every techstack_* key whose value is a list of {name, [logo,] description} dicts,
    find its {{KEY}} placeholder paragraph and replace it with a styled 2-col table.
    Returns the set of keys that were rendered as tables (so the caller skips the
    optional-section drop for them — their placeholder paragraph is already gone and
    a table now sits under the heading)."""
    rendered = set()
    for key, val in list(replacements.items()):
        if not key.startswith("techstack_") or not isinstance(val, list) or not val:
            continue
        if not all(isinstance(it, dict) and "name" in it and "description" in it for it in val):
            continue
        placeholder = "{{" + key.upper() + "}}"
        target = None
        for p in doc.paragraphs:
            if placeholder in (p.text or ""):
                target = p
                break
        if target is None:
            continue
        # Build the table at the end of the document, then move it into position
        tbl = _build_techstack_table(doc, val)
        # Move table to just after the target paragraph
        target._p.addnext(tbl._element)
        # Remove the placeholder paragraph (its content is now in the table)
        target._p.getparent().remove(target._p)
        # Mark replacement so the text-replacement step doesn't try to fill it
        replacements[key] = ""
        rendered.add(key)
    return rendered


def normalize_body_spacing(doc) -> int:
    """Clamp every BODY paragraph (non-heading, non-caption, non-image) to a
    consistent rhythm: SB=0pt, SA=6pt. The template carries a grab-bag of
    paragraph styles inherited from PMO34129 — some sit at SA=20pt, others
    at SA=0 — producing the "chỗ thì trống nhiều chỗ thì gần" defect.

    PRESERVES:
      - The first bullet after a Figure caption (it explicitly carries
        SB=12pt as breathing room between caption table and bullet text).
      - Any paragraph with an intentionally-set SB ≥ 8pt (likely a
        first-after-something marker we put there on purpose).
    """
    touched = 0
    BODY_SB = Pt(0)
    BODY_SA = Pt(6)
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if sname.startswith("Heading") or sname in ("Title", "Subtitle", "Caption"):
            continue
        # Skip image-bearing paragraphs (they own SB=12pt SA=6pt for the figure rhythm)
        if p._p.find(".//" + qn("w:drawing")) is not None:
            continue
        # Skip empty paragraphs (used as page-break or layout spacers)
        if not (p.text or "").strip():
            continue
        pf = p.paragraph_format
        sb = pf.space_before.pt if pf.space_before else 0
        sa = pf.space_after.pt if pf.space_after else 0
        # Preserve intentional larger SB (we explicitly set 12pt for the first
        # bullet after a Figure caption — see _build_bullet_paragraph).
        if sb >= 8:
            # Only clamp SA if it's wildly different from 6pt
            if abs(sa - 6) > 0.1:
                pf.space_after = BODY_SA
                touched += 1
            continue
        if abs(sb - 0) > 0.1:
            pf.space_before = BODY_SB
            touched += 1
        if abs(sa - 6) > 0.1:
            pf.space_after = BODY_SA
            touched += 1
    return touched


def center_image_paragraphs(doc) -> int:
    """Every paragraph that contains a w:drawing (an inline image) gets
    centered. The proposal template ships several diagrams (3.5.3 Testing
    Process, Branch Model, etc.) at left alignment by default — user wants
    every image visually centered on the page so the captions sit cleanly
    beneath."""
    touched = 0
    for p in doc.paragraphs:
        if p._p.find(".//" + qn("w:drawing")) is not None:
            if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                touched += 1
    return touched


def remove_empty_caption_tables(doc) -> int:
    """Remove 1x1 caption tables left over from the template (legacy artifact)."""
    removed = 0
    for tbl in list(doc.tables):
        rows = tbl.rows
        if len(rows) == 1 and len(rows[0].cells) == 1:
            cell = rows[0].cells[0]
            text = "".join(p.text for p in cell.paragraphs).strip()
            if not text:
                tbl._element.getparent().remove(tbl._element)
                removed += 1
    return removed


def _normalise_subheading(text: str) -> str:
    """Strip leading section numbers and normalise for comparison."""
    t = re.sub(r"^\s*\d+(\.\d+)*\s*[.:)-]?\s*", "", text or "")
    return t.strip().lower()


def _collect_subheadings(diagrams: list, existing_anchor_text: str = "") -> list[str]:
    """Build the H6 sub-heading list from the diagram inventory, skipping any
    diagram whose title matches the existing anchor heading (e.g. the
    template already carries '2.1.1 System Context' as the anchor, so a
    Phase 2 inventory that also includes 'System Context' should not inject
    a duplicate H6 — the image just attaches to 2.1.1)."""
    anchor_norm = _normalise_subheading(existing_anchor_text)
    out: list[str] = []
    for d in diagrams:
        title = d.get("subheading") or d.get("target_heading") or d.get("slug")
        if not title or not str(title).strip():
            continue
        title = str(title).strip()
        if anchor_norm:
            t_norm = _normalise_subheading(title)
            if t_norm == anchor_norm or t_norm in anchor_norm or anchor_norm in t_norm:
                continue  # already covered by the existing anchor heading
        out.append(title)
    return out


def build(template: Path, replacements: dict, diagrams: list, out: Path) -> None:
    out.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(template, out)

    doc = Document(str(out))

    # 1. Text replacements across body, tables, and headers/footers.
    # Pre-pass: any techstack_* value that's a list of {name, description} dicts
    # is converted to a styled 2-col table at the placeholder position.
    rendered_tech = render_techstack_tables(doc, replacements)
    if rendered_tech:
        print(f"Techstack tables rendered: {len(rendered_tech)} ({', '.join(sorted(rendered_tech))})")

    total = 0
    for p in doc.paragraphs:
        total += replace_in_paragraph(p, replacements)
    for t in doc.tables:
        total += replace_in_table(t, replacements)
    for section in doc.sections:
        total += replace_in_section_parts(section, replacements)
    print(f"Text replacements: {total}")

    # 2. Inject per-project H6 sub-headings under System Overview.
    #    Skip any diagram whose title already matches the existing anchor
    #    (avoids "2.1.1 System Context" / "2.1.2 System Context" duplicates).
    subs = _collect_subheadings(diagrams, existing_anchor_text=SUBHEADING_ANCHOR)
    injected = inject_subheadings_after_anchor(doc, SUBHEADING_ANCHOR, subs)
    print(f"Sub-headings injected after '{SUBHEADING_ANCHOR}': {injected}")

    # 3. Drop optional empty sections (template defines them; we strip if data is null).
    # The replacements key used for each is derived from heading text by
    # lowercasing + underscoring (see drop_section_if_empty). The headings
    # below have been chosen so the derived key matches the documented
    # Phase 4 schema (mobile_app_strategy, techstack_data, techstack_ai).
    # Heading text -> explicit content key. The Data/AI sub-headings now carry a
    # fixed label ("Data" / "AI") with the {{TECHSTACK_*}} placeholder in the body
    # paragraph beneath, so the content key must be passed explicitly. Skip the
    # drop for any techstack section already rendered as a table (its placeholder
    # is gone and a table sits under the heading — dropping would delete the table).
    drop_specs = [("Mobile App Strategy", None)]
    if "techstack_data" not in rendered_tech:
        drop_specs.append(("Data", "techstack_data"))
    if "techstack_ai" not in rendered_tech:
        drop_specs.append(("AI", "techstack_ai"))
    dropped = 0
    for heading_text, ckey in drop_specs:
        if drop_section_if_empty(doc, heading_text, replacements, ckey):
            dropped += 1
    print(f"Optional sections dropped: {dropped}")

    # 4. Diagram insertions.
    inserted = 0
    for d in diagrams:
        png = Path(d["png"])
        if not png.exists():
            print(f"  ! diagram PNG missing: {png}")
            continue
        ok = insert_image_after_heading(
            doc,
            d["target_heading"],
            png,
            d.get("caption", ""),
            width_inches=d.get("width_inches", 6.5),
            intro_paragraph=d.get("intro_paragraph"),
            explanation_bullets=d.get("explanation_bullets"),
        )
        marker = "ok" if ok else "skipped (heading not found)"
        n_bullets = len(d.get("explanation_bullets") or [])
        print(f"  [{marker}] {d['slug']} -> {d['target_heading']} ({n_bullets} bullets)")
        if ok:
            inserted += 1
    print(f"Diagrams inserted: {inserted}/{len(diagrams)}")

    # 5. Cleanup empty 1x1 caption tables.
    removed_tables = remove_empty_caption_tables(doc)
    print(f"Empty caption tables removed: {removed_tables}")

    # 6a. Pull heading text back to flush-left (template inherits L=36 which
    #     makes headings appear deeper-indented than their body content).
    n_norm = normalize_heading_indents(doc)
    print(f"Heading indents normalized: {n_norm}")

    # 6b. Bump heading space-before/after so section transitions are visible.
    n_sp = boost_heading_spacing(doc)
    print(f"Heading spacing boosted: {n_sp}")

    # 6c. Every Figure caption table needs whitespace after it — body text
    #     touching the caption is a layout defect the user has flagged.
    n_cap = ensure_caption_followed_by_gap(doc, min_pt=12)
    print(f"Paragraphs given gap-after-caption: {n_cap}")

    # 6d. Center every inline image (template-inherited diagrams default to left).
    n_img = center_image_paragraphs(doc)
    if n_img:
        print(f"Image paragraphs centered: {n_img}")

    # 6e. (Disabled) Body paragraph spacing normalization — removed
    #     because the broad sweep regressed the caption-to-bullet gap.
    #     Per [[feedback-dont-touch-fixed]], only touch what's broken.

    # 6f. Justify body.
    justified = justify_body(doc)
    print(f"Paragraphs justified: {justified}")

    # 7. Save.
    doc.save(str(out))

    # 8. Final XML-level placeholder sweep — catches TOC cache, comments etc.
    n_xml = replace_remaining_placeholders_xml(out, replacements)
    if n_xml:
        print(f"XML-level placeholder fills: {n_xml}")

    # 9. autoHyphenation + updateFields patch on the saved file.
    n_num = patch_numbering_xml(out)
    if n_num:
        print(f"Heading numbering left-aligned: {n_num} level(s)")
    n_fig = renumber_figure_sequence_cache(out)
    if n_fig:
        print(f"Figure SEQ cache renumbered: {n_fig}")
    patch_settings_xml(out)

    print(f"\nOutput: {out}  ({out.stat().st_size:,} bytes)")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", type=Path, required=True)
    ap.add_argument("--replacements", type=Path, required=True)
    ap.add_argument("--diagrams", type=Path, required=True)
    ap.add_argument("--out", type=Path, required=True)
    args = ap.parse_args()

    if not args.template.exists():
        print(f"! template missing: {args.template}", file=sys.stderr); return 1
    if not args.replacements.exists():
        print(f"! replacements missing: {args.replacements}", file=sys.stderr); return 1
    if not args.diagrams.exists():
        print(f"! diagrams missing: {args.diagrams}", file=sys.stderr); return 1

    replacements = load_json(args.replacements)
    diagrams = load_json(args.diagrams)
    build(args.template, replacements, diagrams, args.out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
