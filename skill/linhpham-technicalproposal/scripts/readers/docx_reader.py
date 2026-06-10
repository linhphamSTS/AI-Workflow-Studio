#!/usr/bin/env python3
"""Extract text from a .docx file. Tables are flattened into TSV rows."""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document


def read(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []

    body = doc.element.body
    for child in body:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            # python-docx paragraph
            for p in doc.paragraphs:
                if p._p is child:
                    style = p.style.name if p.style else ""
                    text = p.text.strip()
                    if text:
                        if style.startswith("Heading"):
                            lines.append(f"\n# {text}")
                        else:
                            lines.append(text)
                    break
        elif tag == "tbl":
            for t in doc.tables:
                if t._tbl is child:
                    for row in t.rows:
                        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                        lines.append("\t".join(cells))
                    lines.append("")
                    break
    return "\n".join(lines)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("path", type=Path)
    args = ap.parse_args()
    if not args.path.exists():
        print(f"! file not found: {args.path}", file=sys.stderr)
        return 1
    sys.stdout.write(read(args.path))
    return 0


if __name__ == "__main__":
    sys.exit(main())
