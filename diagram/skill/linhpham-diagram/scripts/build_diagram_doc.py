#!/usr/bin/env python3
"""Per-diagram DESCRIPTION Word doc (.docx) — one document per diagram, laid out
exactly like a figure block in a professional technical proposal:

    <diagram name>              (heading)
    <intro paragraph>           (one justified context paragraph ABOVE the image)
    <image>                     (centered)
    Figure N: <caption>         (caption line BELOW the image)
    ● <Bold name> — description  (explanation bullets BELOW the caption)
    ● …

That is the whole document. It deliberately does NOT include "how it was
generated" / "how to reproduce" / engine notes — a proposal figure never shows
those; the reader wants the picture and a description of what is in it. Each
explanation bullet is self-contained: it names a component in bold and folds in
what the component is AND how it connects (mechanism), matching the proposal
convention "**Component**: role; talks to X over HTTPS, publishes events to Y".
The separator after the bold label is a COLON, never a dash: a spaced em-dash is
one of the clearest tells of machine-written text and is rejected on sight.

The content is DERIVED FROM THE SPEC that produced the diagram, so it can never
drift from the picture.

Two ways to describe a diagram:
  • describe(spec, kind, slug)  — auto-derive from a cloud / graph / sequence spec
  • pass a ready-made descriptor (e.g. Phase 3's diagrams.json entry, which already
    carries caption / intro_paragraph / explanation_bullets)

CLI:
  python build_diagram_doc.py --spec s.json --kind cloud --png x.png --out x.docx
  python build_diagram_doc.py --meta descriptor.json --png x.png --out x.docx
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from collections import defaultdict
from pathlib import Path

from PIL import Image


# ---- describe: derive a proposal-style figure block from the generating spec ----
def _icon_human(ref: str) -> str:
    if not ref or "/" not in ref:
        return ref or ""
    prov, stem = ref.split("/", 1)
    stem = stem.split("/")[-1]
    return f"{stem.replace('-', ' ')} ({prov.upper()})"


def _caption(dtype: str, title: str) -> str:
    """A proposal caption reads '<Type> — <Scope>'. Most titles already do; if a
    title has no dash, prefix the diagram type."""
    t = (title or dtype).strip()
    return t if ("—" in t or " - " in t) else f"{dtype} — {t}"


def _cloud_type(spec, slug):
    t = (spec.get("title") or slug or "").lower()
    if "kubernetes" in t or slug == "k8s_topology":
        return "Kubernetes Topology"
    if "docker" in t:
        return "Container Composition"
    if "ci/cd" in t or slug == "cicd":
        return "CI/CD Pipeline"
    if "gitops" in t:
        return "GitOps Delivery Flow"
    if "data pipeline" in t:
        return "Data Pipeline"
    if "microservice" in t:
        return "Microservices Decomposition"
    if "container diagram" in t or slug == "c4_container":
        return "C4 Container Diagram"
    if "deployment" in t:
        return "UML Deployment Diagram"
    if "hybrid" in t or "on-prem" in t:
        return "Hybrid / On-Prem Architecture"
    if "rag" in t or "model" in t:
        return "AI / RAG Architecture"
    return "Cloud Reference Architecture"


_ROLE_KW = {
    "edge": ("edge", "global", "dns", "cdn", "waf", "front", "armor", "cloudfront", "route 53",
             "route53", "perimeter", "dmz", "internet"),
    "lb": ("load balanc", "app gateway", "application gateway", "global lb", "alb", "elb", "ingress"),
    "app": ("eks", "aks", "gke", "kubernetes", "cloud run", "compute", " app ", "app ", "service",
            "pool", "deployment", "api", "backend", "worker", "function", "fargate", "container"),
    "data": ("data", "sql", "aurora", "database", " db", "cache", "redis", "elasticache",
             "memorystore", "storage", "blob", "s3", "lake", "warehouse", "redshift", "opensearch",
             "endpoint", "persist"),
    "msg": ("messag", "queue", "kafka", "msk", "pubsub", "pub/sub", "service bus", "sqs", "sns",
            "event", "stream", "rabbit", "amqp"),
}


def _classify_col(col):
    """Best-effort role bucket for a tier column, from its boundary label + node labels."""
    text = " ".join([(col.get("boundary") or {}).get("label", "")]
                    + [n.get("label", "") for n in col.get("nodes", [])]).lower()
    for role in ("msg", "data", "lb", "app", "edge"):   # most-specific first
        if any(k in text for k in _ROLE_KW[role]):
            return role
    return "other"


_DTYPE_KEEP = {"CI/CD", "GitOps", "Kubernetes", "AI", "RAG", "AWS", "Azure", "GCP", "UML", "C4"}


def _dtype_phrase(dtype):
    """Lower-case a diagram-type name for mid-sentence use, but keep acronyms and
    proper nouns capitalised ('CI/CD pipeline', 'GitOps delivery flow')."""
    return " ".join(w if w in _DTYPE_KEEP else w.lower() for w in dtype.split())


def _names(nodes, cap=5):
    """Join node/label dicts into a readable 'A, B and C' phrase. Accepts dicts with
    'label' and/or 'id' (never eager-indexes 'id')."""
    labs = [(n.get("label") or n.get("id") or "").split("(")[0].strip() for n in nodes]
    labs = [l for l in labs if l]
    if not labs:
        return ""
    if len(labs) > cap:
        return ", ".join(labs[:cap]) + ", …"
    if len(labs) == 1:
        return labs[0]
    return ", ".join(labs[:-1]) + " and " + labs[-1]


def _col_labels(spec):
    """Human labels for the columns, in order. Prefer the boundary label; else the
    first node's label; never a raw column id (e.g. 'cl')."""
    out = []
    for c in spec.get("columns", []):
        b = (c.get("boundary") or {}).get("label", "")
        if not b:
            nodes = c.get("nodes", [])
            b = (nodes[0].get("label") or nodes[0].get("id", "")) if nodes else ""
        b = b.split("(")[0].strip()
        if b:
            out.append(b)
    return out


def _narrate_cloud(spec, dtype):
    """A senior-SA-style intro that fits the diagram FAMILY — never 'reads left-to-right
    across N tiers' and never the sync/async legend (that is in the image itself)."""
    cols = spec.get("columns", [])
    labels = _col_labels(spec)
    shared = spec.get("shared")
    wraps = spec.get("wraps", []) or []

    # 1) STAGE PIPELINES — describe the left-to-right stage flow
    if dtype in ("CI/CD Pipeline", "GitOps Delivery Flow", "Data Pipeline") and len(labels) >= 2:
        return (f"A {_dtype_phrase(dtype)} that runs left to right through {' → '.join(labels)}, "
                "carrying a change from the first stage to the last.")

    # 2) MICROSERVICES — bounded contexts integrated over sync + async
    if dtype == "Microservices Decomposition" and labels:
        ctxs = [l for l in labels if "context" in l.lower()] or labels
        return (f"A microservices decomposition into {_names([{'label': l} for l in ctxs])}, "
                "each context owning its own service and datastore and integrating with the others "
                "over synchronous APIs and asynchronous events.")

    # 3) RAG — orchestration → models → retrieval
    if dtype == "AI / RAG Architecture":
        return ("A retrieval-augmented AI architecture: a user request is orchestrated, dispatched to "
                "the model providers, and grounded with context retrieved from the vector store before "
                "the answer is returned.")

    # 4) REFERENCE ARCHITECTURES (cloud / k8s / on-prem / container) — trace the request path
    buckets = {"edge": [], "lb": [], "app": [], "data": [], "msg": [], "other": []}
    for c in cols:
        buckets[_classify_col(c)].extend(c.get("nodes", []))
    clauses = []
    entry = buckets["edge"] + buckets["lb"]
    if entry:
        clauses.append(f"inbound traffic terminates at {_names(entry)}")
    if buckets["app"]:
        clauses.append(f"{'is served by' if clauses else 'the workload runs on'} {_names(buckets['app'])}")
    if buckets["data"]:
        clauses.append(f"state is held in {_names(buckets['data'])}")
    lead = f"A {_dtype_phrase(dtype)} in which"
    if clauses:
        body = clauses[0] + "".join((f", and {c}" if i == len(clauses) - 1 else f", {c}")
                                    for i, c in enumerate(clauses[1:], 1))
        sentences = [f"{lead} {body}."]
    else:
        sentences = [f"A {_dtype_phrase(dtype)} spanning {_names([{'label': l} for l in labels]) or 'its tiers'}."]
    if buckets["msg"]:
        sentences.append(f"Asynchronous, event-driven work flows through {_names(buckets['msg'])}.")
    tail = ""
    if wraps:
        tail = f"The estate sits inside {_names([{'label': w.get('label', '')} for w in wraps])}"
    if shared:
        sh = _names(shared.get("nodes", []))
        band = shared.get("label", "shared services").split("(")[0].strip()
        tail = (f"{tail}, with {band.lower()} ({sh}) applied across every tier" if tail
                else f"{band} ({sh}) apply across every tier")
    if tail:
        sentences.append(tail + ".")
    return " ".join(sentences)


def _cloud_connections(spec):
    """Map node id -> list of human connection phrases folded into its bullet."""
    label = {}
    for c in spec.get("columns", []):
        for n in c.get("nodes", []):
            label[n["id"]] = n.get("label", n["id"])
    if spec.get("shared"):
        for n in spec["shared"].get("nodes", []):
            label[n["id"]] = n.get("label", n["id"])
    outs = defaultdict(list)
    for e in spec.get("edges", []):
        dashed = e.get("style") == "dashed"
        tgt = label.get(e.get("to"), e.get("to"))
        lbl = e.get("label")
        if lbl:
            verb = "publishes" if dashed else "sends"
            phrase = f"{verb} {lbl} to {tgt}" if dashed else f"{lbl} to {tgt}"
        else:
            phrase = f"{'streams events to' if dashed else 'connects to'} {tgt}"
        outs[e.get("from")].append(phrase)
    return outs


def _describe_cloud(spec, slug):
    dtype = _cloud_type(spec, slug)
    cols = spec.get("columns", [])
    wraps = spec.get("wraps", []) or []
    shared = spec.get("shared")
    intro = _narrate_cloud(spec, dtype)

    outs = _cloud_connections(spec)
    bullets = []
    for w in wraps:
        style = "transparent, dashed network boundary" if w.get("dashed") else "network boundary"
        badge = f", badged {_icon_human(w['icon'])}" if w.get("icon") else ""
        bullets.append((w.get("label", "Boundary"),
                        f"{style}{badge} spanning the tiers it encloses"))
    for c in cols:
        b = c.get("boundary")
        for n in c.get("nodes", []):
            what = _icon_human(n.get("icon", "")) or "component"
            if n.get("tech"):
                what += f", {n['tech']}"
            where = f" in {b['label'].split('(')[0].strip()}" if b else ""
            conns = outs.get(n["id"], [])
            tail = f"; {', '.join(conns)}" if conns else ""
            bullets.append((n.get("label", n["id"]), f"{what}{where}{tail}"))
    if shared:
        names = ", ".join(n.get("label", n["id"]) for n in shared.get("nodes", []))
        anchor = shared.get("anchor") or {}
        lbl = anchor.get("label", "")
        bullets.append((shared.get("label", "Shared Services").split("(")[0].strip(),
                        f"cross-cutting services ({names}) applied to every tier"
                        + (f" — {lbl}" if lbl else "")))
    return {"title": spec.get("title", slug), "subheading": dtype,
            "caption": _caption(dtype, spec.get("title", slug)),
            "intro": intro, "bullets": bullets}


_GRAPH_TYPE = {
    "c4": "C4 Diagram", "class": "UML Class Diagram", "erd": "Entity-Relationship Diagram",
    "dfd": "Data-Flow Diagram", "bpmn": "BPMN Process", "swimlane": "Swimlane Process",
    "flowchart": "Flowchart", "state": "State Machine", "orgchart": "Org Chart",
    "mindmap": "Mind Map", "network": "Network Topology",
}


def _describe_graph(spec, slug):
    dt = spec.get("diagram_type", "flowchart")
    dtype = _GRAPH_TYPE.get(dt, dt.replace("_", " ").title())
    direction = spec.get("direction", "TB")
    label = {n["id"]: n.get("label", n["id"]) for n in spec.get("nodes", [])}
    outs = defaultdict(list)
    for e in spec.get("edges", []):
        tgt = label.get(e.get("to"), e.get("to"))
        lbl = f" ({e['label']})" if e.get("label") else ""
        outs[e.get("from")].append(f"{tgt}{lbl}")
    nnodes, nedges = len(spec.get("nodes", [])), len(spec.get("edges", []))
    unit = {"c4": "elements", "class": "classes", "erd": "entities", "dfd": "processes and stores",
            "bpmn": "activities", "swimlane": "activities", "flowchart": "steps", "state": "states",
            "orgchart": "roles", "mindmap": "branches", "network": "nodes"}.get(dt, "elements")
    key = _names(spec.get("nodes", []), cap=4)
    process = dt in ("flowchart", "bpmn", "swimlane", "state", "dfd")
    if process:
        intro = (f"{spec.get('title', slug)} traces the flow through {nnodes} {unit}"
                 + (f" — {key} —" if key else "")
                 + f" and the {nedges} transition{'s' if nedges != 1 else ''} between them.")
    else:
        intro = (f"{spec.get('title', slug)} models {nnodes} {unit}"
                 + (f" — {key} —" if key else "")
                 + f" and the {nedges} relationship{'s' if nedges != 1 else ''} that connect them.")
    bullets = []
    for n in spec.get("nodes", []):
        bits = [str(n[k]) for k in ("role", "type", "tech", "subtitle") if n.get(k)]
        for k in ("fields", "attributes", "methods"):
            if n.get(k):
                bits.append(f"{k}: " + ", ".join(n[k]))
        conns = outs.get(n["id"], [])
        if conns:
            bits.append("connects to " + ", ".join(conns))
        bullets.append((n.get("label", n["id"]), "; ".join(bits) if bits else "node"))
    for cl in spec.get("clusters", []) or []:
        bullets.append((cl.get("label", cl["id"]),
                        "grouping: " + ", ".join(cl.get("members", []))))
    return {"title": spec.get("title", slug), "subheading": dtype,
            "caption": _caption(dtype, spec.get("title", slug)),
            "intro": intro, "bullets": bullets}


def _describe_sequence(spec, slug):
    parts = spec.get("participants", [])
    msgs = spec.get("messages", [])
    frags = spec.get("fragments", []) or []
    intro = (f"{spec.get('title', slug)} traces how {_names(parts)} collaborate, exchanging "
             f"{len(msgs)} message{'s' if len(msgs) != 1 else ''} in time order"
             + (f", including {len(frags)} conditional or repeated stretch"
                f"{'es' if len(frags) != 1 else ''}" if frags else "")
             + ".")
    label = {p["id"]: p.get("label", p["id"]) for p in parts}
    bullets = [(p.get("label", p["id"]), p.get("role", "participant")) for p in parts]
    for i, m in enumerate(msgs, 1):
        src = label.get(m.get("from"), m.get("from")); tgt = label.get(m.get("to"), m.get("to"))
        bullets.append((f"Step {i}",
                        f"{src} → {tgt}: {m.get('label','')} ({m.get('kind','sync')})"))
    return {"title": spec.get("title", slug), "subheading": "UML Sequence Diagram",
            "caption": _caption("Sequence Diagram", spec.get("title", slug)),
            "intro": intro, "bullets": bullets}


def describe(spec: dict, kind: str, slug: str = "") -> dict:
    slug = slug or spec.get("slug", "")
    if kind == "cloud":
        return _describe_cloud(spec, slug)
    if kind == "sequence":
        return _describe_sequence(spec, slug)
    return _describe_graph(spec, slug)


# ---- normalise a descriptor (accept our shape OR a diagrams.json entry) ---------
def _normalise(desc: dict) -> dict:
    """Accept either this module's shape (subheading/caption/intro/bullets) or a
    Phase-3 diagrams.json entry (subheading/caption/intro_paragraph/
    explanation_bullets, bullets as '**Name** — desc' strings)."""
    out = {
        "title": desc.get("title") or desc.get("subheading") or "Diagram",
        "subheading": desc.get("subheading") or desc.get("title") or "Diagram",
        "caption": desc.get("caption") or desc.get("title") or "Diagram",
        "intro": desc.get("intro") or desc.get("intro_paragraph") or "",
    }
    bullets = desc.get("bullets")
    if bullets is None:
        # diagrams.json style: list of "**Name** — description" strings
        bullets = []
        for s in desc.get("explanation_bullets", []) or []:
            bullets.append(_split_markdown_bullet(str(s)))
    out["bullets"] = [b if isinstance(b, (list, tuple)) else (str(b), "") for b in bullets]
    return out


def _split_markdown_bullet(s: str):
    """'**Name** — desc' -> ('Name', 'desc'); tolerant of missing bold / dash."""
    s = s.strip()
    if s.startswith("**") and "**" in s[2:]:
        name, rest = s[2:].split("**", 1)
        rest = rest.lstrip()
        for sep in ("—", " - ", ":"):
            if rest.startswith(sep):
                rest = rest[len(sep):].strip(); break
        return (name.strip(), rest)
    for sep in ("—", " - "):
        if sep in s:
            a, b = s.split(sep, 1)
            return (a.strip(" *"), b.strip())
    return (s, "")


# ---- render the .docx -----------------------------------------------------------
def _fit_inches(png: Path):
    """Return (width_in, height_in) fitting a Letter text column (6.5in x 8.2in)."""
    with Image.open(png) as im:
        w, h = im.size
    aspect = w / h if h else 1.0
    wi, hi = 6.5, 6.5 / aspect
    if hi > 8.2:
        hi, wi = 8.2, 8.2 * aspect
    return wi, hi


def build_doc(desc: dict, png: Path, out: Path):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = _normalise(desc)
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(1.0)

    # heading (the diagram's short name)
    doc.add_heading(d["subheading"], level=1)

    # intro paragraph — one justified context sentence ABOVE the image
    if d["intro"]:
        p = doc.add_paragraph(d["intro"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)

    # image (centered)
    if png and Path(png).exists():
        wi, hi = _fit_inches(Path(png))
        pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.paragraph_format.space_before = Pt(6); pic.paragraph_format.space_after = Pt(4)
        pic.add_run().add_picture(str(png), width=Inches(wi))

    # "Figure 1: <caption>" — caption line BELOW the image
    cap = re.sub(r"^\s*Figure\s+\d+\s*[:.]?\s*", "", d["caption"]).strip() or d["caption"]
    capp = doc.add_paragraph(); capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    capp.paragraph_format.space_after = Pt(8)
    cr = capp.add_run("Figure 1: "); cr.bold = True; cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(0x60, 0x6a, 0x78)
    cr2 = capp.add_run(cap); cr2.italic = True; cr2.font.size = Pt(9)
    cr2.font.color.rgb = RGBColor(0x60, 0x6a, 0x78)

    # explanation bullets: "● **Name**: description".
    # A colon, not a dash: a spaced em-dash is one of the clearest tells of
    # machine-written text and the client rejects it on sight.
    for name, body in d["bullets"]:
        para = doc.add_paragraph(style="List Bullet")
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(4)
        rn = para.add_run(str(name)); rn.bold = True
        if body:
            para.add_run(f": {body}")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--spec", type=Path, help="a cloud/graph/sequence spec JSON")
    ap.add_argument("--kind", choices=["cloud", "graph", "sequence"], default="graph")
    ap.add_argument("--meta", type=Path, help="a ready-made descriptor JSON (overrides --spec)")
    ap.add_argument("--png", type=Path, required=True)
    ap.add_argument("--out", type=Path, required=True)
    a = ap.parse_args()
    if a.meta:
        desc = json.loads(a.meta.read_text(encoding="utf-8"))
    elif a.spec:
        desc = describe(json.loads(a.spec.read_text(encoding="utf-8")), a.kind)
    else:
        print("need --spec or --meta", file=sys.stderr); sys.exit(2)
    build_doc(desc, a.png, a.out)
    print(f"Wrote {a.out}")


if __name__ == "__main__":
    main()
