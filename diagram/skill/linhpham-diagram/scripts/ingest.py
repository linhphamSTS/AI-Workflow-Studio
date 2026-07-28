#!/usr/bin/env python3
"""Ingest a project folder into ONE text digest the diagram skill can analyse.

Extracts readable text from .txt / .md / .markdown / .csv / .json / .yaml / .docx /
.xlsx / .xlsm / .pdf under a folder (recursively, skipping noise dirs), and writes a
consolidated Markdown digest with one section per file. Phase 1 (Refine) reads the
digest to understand the system/requirements and propose the RIGHT SET of diagrams —
exactly the "point at a folder, ingest the docs, refine into diagram prompts" flow the
technical-proposal skill uses.

Readers: .docx → python-docx (paragraphs + tables); .xlsx/.xlsm → openpyxl (sheet
rows); .pdf → PyMuPDF (fitz) with a PyPDF2 fallback; text-like → read directly.

Usage:
  python ingest.py --dir <folder> [--out <digest.md>] [--max-chars-per-file N]
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

try:  # so file names with non-ASCII (e.g. Vietnamese) never crash the cp1252 console
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:  # noqa: BLE001
    pass

SKIP_DIRS = {".git", "node_modules", "__pycache__", ".venv", "venv", "output", ".idea",
             ".vscode", "dist", "build", ".next", "bin", "obj", ".gradle", "target"}
TEXT_EXT = {".txt", ".md", ".markdown", ".csv", ".log", ".json", ".yaml", ".yml"}
DOC_EXT = {".docx"}
XLS_EXT = {".xlsx", ".xlsm"}
PDF_EXT = {".pdf"}
SUPPORTED = TEXT_EXT | DOC_EXT | XLS_EXT | PDF_EXT


def _read_text(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="replace")


def _read_docx(p: Path) -> str:
    from docx import Document
    doc = Document(str(p))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_xlsx(p: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"### Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(" | ".join(cells))
    wb.close()
    return "\n".join(out)


def _read_pdf(p: Path) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(p))
        out = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(out)
    except Exception:  # noqa: BLE001 — fall back to PyPDF2
        from PyPDF2 import PdfReader
        r = PdfReader(str(p))
        return "\n".join((pg.extract_text() or "") for pg in r.pages)


def extract(p: Path) -> str:
    ext = p.suffix.lower()
    if ext in DOC_EXT:
        return _read_docx(p)
    if ext in XLS_EXT:
        return _read_xlsx(p)
    if ext in PDF_EXT:
        return _read_pdf(p)
    if ext in TEXT_EXT:
        return _read_text(p)
    return ""


def ingest(root: Path, out: Path, max_chars: int = 20000):
    files = []
    for p in sorted(root.rglob("*")):
        if p.is_dir() or any(part in SKIP_DIRS for part in p.parts):
            continue
        if p.suffix.lower() in SUPPORTED:
            files.append(p)
    sections = [f"# Ingested digest — {root.name}",
                f"_{len(files)} file(s) ingested from `{root}`._", ""]
    stats = []
    for p in files:
        rel = p.relative_to(root)
        try:
            text = extract(p).strip()
        except Exception as e:  # noqa: BLE001 — never fail the whole ingest on one file
            sections.append(f"## {rel}\n\n_(could not read: {type(e).__name__}: {e})_\n")
            stats.append((str(rel), "ERR")); continue
        trunc = ""
        if len(text) > max_chars:
            dropped = len(text) - max_chars
            text = text[:max_chars]
            # Say WHAT was lost and where to get it — a silent "…(truncated)" invites the
            # reader to treat a partial document as the whole requirement set.
            trunc = (f"  ⚠ TRUNCATED at {max_chars:,} chars — {dropped:,} more remain; "
                     f"open the original file for the rest")
        sections.append(f"## {rel}{trunc}\n\n{text or '_(no extractable text)_'}\n")
        stats.append((str(rel), f"{len(text)} chars"))
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(sections), encoding="utf-8")
    return files, stats, out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--dir", type=Path, required=True, help="project folder to ingest")
    ap.add_argument("--out", type=Path, help="digest output path (default: <dir>/output/diagrams/_ingest_digest.md)")
    ap.add_argument("--max-chars-per-file", type=int, default=20000)
    a = ap.parse_args()
    if not a.dir.is_dir():
        print(f"! not a folder: {a.dir}", file=sys.stderr); sys.exit(2)
    out = a.out or (a.dir / "output" / "diagrams" / "_ingest_digest.md")
    files, stats, out = ingest(a.dir, out, a.max_chars_per_file)
    print(f"Ingested {len(files)} file(s) -> {out}")
    for name, st in stats:
        print(f"  - {name}: {st}")
    if not files:
        print("  (no supported files found: .txt .md .csv .json .yaml .docx .xlsx .pdf)")


if __name__ == "__main__":
    main()
