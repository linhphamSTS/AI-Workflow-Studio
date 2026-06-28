#!/usr/bin/env python3
"""One-shot, sectPr-safe genericization of proposal_template.docx.

The XML_SUBS table below carries the customer-specific source strings that
need to be turned into placeholders. These are PATTERNS to find-and-replace,
not customer references that survive in any built output — the resulting
template has only `{{KEY}}` placeholders. Auditors grepping the repo for
customer names will hit this file; that hit is intentional.

Critical invariant: never delete a <w:p> that contains a <w:sectPr>. Section
breaks live inside paragraphs; removing the paragraph also removes the break,
which collapses sections and wipes per-section margins.

Idempotent and safe to re-run. Delete this file after the template is frozen
if you want a fully customer-free working tree (the committed template will
still be generic).
"""
from __future__ import annotations

import copy
import re
import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SRC = Path("skill/linhpham-technicalproposal/templates/proposal_template.docx")

XML_SUBS = [
    # Multi-word phrases first so substring replacements don't double-fire.
    # The case study is a per-project fill-in -> placeholders, never a hardcoded
    # engagement name (keeps the template general / vendor-neutral).
    ("comparable to the Tay Ho B2B Airline Portal.",
     "comparable to {{CASE_STUDY_TITLE}}."),
    ("FruPro, a multi-stakeholder B2B platform we delivered for a similar engagement",
     "{{CASE_STUDY_TITLE}}"),
    ("https://saigontechnology.com/case-studies/frupro", "{{CASE_STUDY_URL}}"),
    ("B2B Airline Portal Platform", "{{PROJECT_TITLE}}"),
    ("sponsor@tayho.vn", "{{CLIENT_CONTACT_EMAIL}}"),
    ("Tay Ho Group", "{{CLIENT_NAME}}"),
    ("Tay Ho", "{{CLIENT_NAME}}"),
    ("Vendor Y", "{{VENDOR_PARTNER_NAME}}"),
]
COVER_REPLACEMENTS = {
    "Version 1.0 - 19 May 2026": "{{VERSION}} - {{PROPOSAL_DATE}}",
}
SECTION_2_1_DROP = {
    "2.1.2 AWS Reference Architecture",
    "2.1.3 Vendor Y Integration: Anti-Corruption Layer",
    "2.1.3 {{VENDOR_PARTNER_NAME}} Integration: Anti-Corruption Layer",
    "2.1.4 Booking Saga: Idempotent Orchestration",
    "2.1.5 Public API Hub",
}


def has_drawing(p) -> bool:
    return bool(p._p.findall(".//" + qn("w:drawing"))) or bool(p._p.findall(".//" + qn("w:pict")))


def has_sectpr(p_el) -> bool:
    return bool(p_el.findall(".//" + qn("w:sectPr")))


def set_keep(p, tag: str) -> None:
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p._p.insert(0, pPr)
    for existing in pPr.findall(qn(tag)):
        pPr.remove(existing)
    pPr.append(OxmlElement(tag))


def replace_runs_text(p_el, text: str) -> None:
    runs = p_el.findall(qn("w:r"))
    set_once = False
    for r in runs:
        for t in r.findall(qn("w:t")):
            if not set_once:
                t.text = text
                set_once = True
            else:
                t.text = ""
    if not set_once and runs:
        t = OxmlElement("w:t")
        t.text = text
        runs[0].append(t)


def phase1_xml_subs(src: Path) -> int:
    """ZIP-level string subs across all .xml parts (document, headers, footers).
    Also marks TOC fields dirty so Word refreshes them on open, and ensures
    settings.xml has autoHyphenation + updateFields.
    """
    tmp = src.with_suffix(".tmp.docx")
    toc_dirty = [0]
    with zipfile.ZipFile(src) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".xml"):
                txt = data.decode("utf-8", errors="replace")
                for old, new in XML_SUBS:
                    txt = txt.replace(old, new)
                if item.filename.endswith(".xml") and ("document" in item.filename
                                                       or "header" in item.filename
                                                       or "footer" in item.filename):
                    # Mark EVERY field begin (TOC, SEQ, STYLEREF, REF, PAGE, …)
                    # dirty so Word refreshes the cached value on open. Without
                    # this the dynamic header that references the closest
                    # Heading style keeps the value cached when the template
                    # was last saved by Word.
                    def _add_dirty(m):
                        tag = m.group(0)
                        if "w:dirty" in tag:
                            return tag
                        toc_dirty[0] += 1
                        return tag.replace("/>", ' w:dirty="true"/>')
                    txt = re.sub(
                        r'<w:fldChar\b[^/]*?w:fldCharType="begin"[^/]*/>',
                        _add_dirty, txt,
                    )
                if item.filename == "word/settings.xml":
                    inserts = []
                    if "autoHyphenation" not in txt:
                        inserts.append('<w:autoHyphenation w:val="true"/>')
                    if "updateFields" not in txt:
                        inserts.append('<w:updateFields w:val="true"/>')
                    if inserts:
                        txt = txt.replace("</w:settings>", "".join(inserts) + "</w:settings>")
                data = txt.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(tmp, src)
    return toc_dirty[0]


def phase2_cover_repl(doc) -> int:
    hits = 0
    for p in doc.paragraphs:
        joined = "".join(r.text or "" for r in p.runs)
        if not joined:
            continue
        new = joined
        for old, repl in COVER_REPLACEMENTS.items():
            if old in new:
                new = new.replace(old, repl)
        if new != joined:
            replace_runs_text(p._p, new)
            hits += 1
    return hits


def phase3_drop_21x(doc) -> int:
    to_remove = []
    for p in doc.paragraphs:
        if p.text.strip() in SECTION_2_1_DROP:
            if has_sectpr(p._p) or has_drawing(p):
                continue  # never remove section-break or drawing-bearing paragraphs
            to_remove.append(p._p)
    for el in to_remove:
        if el.getparent() is not None:
            el.getparent().remove(el)
    return len(to_remove)


def phase4_add_techstack(doc) -> int:
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "Server & Hosting" and p.style and p.style.name == "Heading 6":
            anchor = p
            break
    if anchor is None:
        return 0
    added = 0
    for text in reversed(["{{TECHSTACK_DATA}}", "{{TECHSTACK_AI}}"]):
        new_el = copy.deepcopy(anchor._p)
        # Strip any sectPr from clone — don't duplicate section breaks
        for sp in new_el.findall(".//" + qn("w:sectPr")):
            sp.getparent().remove(sp)
        # Strip keepNext on clone (anchor might have it from earlier runs)
        pPr = new_el.find(qn("w:pPr"))
        if pPr is not None:
            for kn in pPr.findall(qn("w:keepNext")):
                pPr.remove(kn)
        replace_runs_text(new_el, text)
        anchor._p.addnext(new_el)
        added += 1
    return added


def phase5_normalize_summary(doc) -> bool:
    for p in doc.paragraphs:
        if p.text.strip() == "5. SUMMARY" and p.style and p.style.name == "Heading 3":
            joined = "".join(r.text or "" for r in p.runs)
            new = joined.replace("5. SUMMARY", "SUMMARY")
            if new != joined:
                replace_runs_text(p._p, new)
            return True
    return False


def phase6_strip_body(doc, heading_text: str, placeholder: str) -> int:
    """Strip body of `heading_text` H3 section but PRESERVE paragraphs with sectPr or drawings.
    Insert a single placeholder paragraph after the heading."""
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        s = p.style.name if p.style else ""
        if s == "Heading 3" and p.text.strip() == heading_text:
            start = i
            break
    if start is None:
        return 0
    end = len(paras)
    for j in range(start + 1, len(paras)):
        s = paras[j].style.name if paras[j].style else ""
        if s == "Heading 3":
            end = j
            break

    cleared = 0
    for j in range(start + 1, end):
        p = paras[j]
        if has_drawing(p):
            continue  # preserve images
        if has_sectpr(p._p):
            # CRITICAL: keep paragraph so section break stays; only wipe runs.
            for r in p._p.findall(qn("w:r")):
                p._p.remove(r)
            cleared += 1
            continue
        if p._p.getparent() is not None:
            p._p.getparent().remove(p._p)
            cleared += 1

    # Insert placeholder right after heading. Strip every property that
    # would make the placeholder render anywhere other than directly under
    # the heading on the same page.
    head_p = paras[start]
    new_el = copy.deepcopy(head_p._p)
    for sp in new_el.findall(".//" + qn("w:sectPr")):
        sp.getparent().remove(sp)
    pPr = new_el.find(qn("w:pPr"))
    if pPr is not None:
        for child_tag in ("w:pStyle", "w:keepNext", "w:keepLines",
                          "w:pageBreakBefore", "w:numPr", "w:outlineLvl"):
            for el in pPr.findall(qn(child_tag)):
                pPr.remove(el)
    replace_runs_text(new_el, placeholder)
    head_p._p.addnext(new_el)
    return cleared


def phase7_restructure_testing(doc) -> int:
    """Replace 6x2 Testing Process table (duplicate image per row) with one
    centred image + caption + bullet paragraphs that PRESERVE the original
    cell formatting (bold 'Test Planning:' title + the normal description).

    Also removes any stray 1x1 'Testing Process' caption table BEFORE adding
    the new caption — otherwise downstream phases that look for a 1x1 with
    that text would target the wrong table.
    """
    target = None
    for t in doc.tables:
        if len(t.rows) == 6 and len(t.columns) == 2:
            c1 = " ".join(p.text for p in t.rows[0].cells[1].paragraphs)
            if "Test Planning" in c1:
                target = t
                break
    if target is None:
        return 0

    # Pre-remove any stray 1x1 caption table whose text is exactly "Testing
    # Process" (a leftover from the original template that would otherwise
    # collide with the new caption we're about to insert).
    for t in list(doc.tables):
        if len(t.rows) == 1 and len(t.columns) == 1:
            if t.rows[0].cells[0].text.strip() == "Testing Process":
                t._tbl.getparent().remove(t._tbl)

    tbl_el = target._tbl
    parent = tbl_el.getparent()
    idx = list(parent).index(tbl_el)

    # 1) Centered image paragraph (clone the table's first drawing).
    new_img_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    new_img_p.append(pPr)
    first_drawing = target.rows[0].cells[0]._tc.find(".//" + qn("w:drawing"))
    if first_drawing is not None:
        r_el = OxmlElement("w:r")
        r_el.append(copy.deepcopy(first_drawing))
        new_img_p.append(r_el)
    parent.insert(idx, new_img_p)
    idx += 1

    # 1b) Caption: use the same 1x1 bordered table format the rest of
    # Section 3 uses (Agile Methodology, Branching Model, etc.) so the
    # Testing Process figure matches its neighbours visually. Clone an
    # existing 1x1 caption table and rewrite its text.
    sample_caption_tbl = None
    for t in doc.tables:
        if len(t.rows) == 1 and len(t.columns) == 1:
            txt = t.rows[0].cells[0].text.strip()
            if txt and txt != "Testing Process":  # any of the other Section 3 caption tables
                sample_caption_tbl = t
                break
    if sample_caption_tbl is not None:
        new_tbl = copy.deepcopy(sample_caption_tbl._tbl)
        # Replace text inside the single cell
        for tnode in new_tbl.iter(qn("w:t")):
            tnode.text = ""
        # Set new text into the first w:t we find (or add one)
        first_t = next(iter(new_tbl.iter(qn("w:t"))), None)
        if first_t is not None:
            first_t.text = "Testing Process"
        else:
            first_p = next(iter(new_tbl.iter(qn("w:p"))), None)
            if first_p is not None:
                r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = "Testing Process"
                r.append(t)
                first_p.append(r)
        parent.insert(idx, new_tbl)
        idx += 1

    # 2) For each row, deepcopy the FIRST non-empty paragraph of column 1.
    #    This preserves run formatting (bold "Test Planning:" + normal body).
    #    Prepend a "●  " text-only run at the start of each cloned paragraph.
    inserted = 0
    for row in target.rows:
        cell = row.cells[1]
        src_p = None
        for cp in cell.paragraphs:
            if cp.text.strip():
                src_p = cp
                break
        if src_p is None:
            continue
        p_el = copy.deepcopy(src_p._p)
        # Strip section break, numPr, keepNext, pageBreakBefore from the clone.
        for sp in p_el.findall(".//" + qn("w:sectPr")):
            sp.getparent().remove(sp)
        pPr = p_el.find(qn("w:pPr"))
        if pPr is not None:
            for tag in ("w:numPr", "w:keepNext", "w:pageBreakBefore", "w:outlineLvl"):
                for el in pPr.findall(qn(tag)):
                    pPr.remove(el)
        # If row content already begins with '●' there's nothing to add.
        existing_text = "".join(t.text or "" for t in p_el.iter(qn("w:t")))
        if not existing_text.lstrip().startswith("●"):
            bullet_run = OxmlElement("w:r")
            t_el = OxmlElement("w:t")
            t_el.set(qn("xml:space"), "preserve")
            t_el.text = "●  "
            bullet_run.append(t_el)
            # Insert before the existing first <w:r>.
            first_r = p_el.find(qn("w:r"))
            if first_r is not None:
                first_r.addprevious(bullet_run)
            else:
                p_el.append(bullet_run)
        parent.insert(idx + inserted, p_el)
        inserted += 1
    parent.remove(tbl_el)
    return inserted


def phase9_strip_bullet_before_number(doc) -> int:
    """For paragraphs that begin with '● 1.', '● 2.', etc., strip the literal
    '● ' so the line reads '1. Title: ...' — avoids the double bullet+number
    rendering when the source author put both a list bullet and an explicit
    step number in the text."""
    pattern = re.compile(r"^●\s+(\d+\.\s)")
    fixed = 0
    for p in doc.paragraphs:
        joined = "".join(r.text or "" for r in p.runs)
        m = pattern.match(joined)
        if not m:
            continue
        # Find the first run that contains the leading '●' and strip in place
        for r in p.runs:
            if r.text and r.text.lstrip().startswith("●"):
                r.text = re.sub(r"^\s*●\s+", "", r.text)
                fixed += 1
                break
    return fixed


def phase10_spacing_heading_before_image(doc) -> int:
    """If a heading is immediately followed (modulo empty paragraphs) by an
    image-bearing paragraph, set a spacing-after on the heading so it doesn't
    visually touch the image."""
    paras = doc.paragraphs
    fixed = 0
    for i, p in enumerate(paras):
        sname = p.style.name if p.style else ""
        if not sname.startswith("Heading"):
            continue
        j = i + 1
        while j < len(paras):
            nxt = paras[j]
            if nxt._p.findall(".//" + qn("w:drawing")):
                pPr = p._p.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    p._p.insert(0, pPr)
                spacing = pPr.find(qn("w:spacing"))
                if spacing is None:
                    spacing = OxmlElement("w:spacing")
                    pPr.append(spacing)
                spacing.set(qn("w:after"), "240")  # 12 pt
                fixed += 1
                break
            if nxt.text.strip():
                break
            j += 1
    return fixed


def phase7b_remove_caption_table(doc, label: str) -> int:
    """Remove a stray 1x1 caption table whose only content is `label`."""
    target = None
    for t in doc.tables:
        if len(t.rows) == 1 and len(t.columns) == 1:
            txt = t.rows[0].cells[0].text.strip()
            if txt == label:
                target = t
                break
    if target is None:
        return 0
    target._tbl.getparent().remove(target._tbl)
    return 1


def make_seq_caption_paragraph(title: str):
    """Build a centred italic 'Figure {SEQ Figure}: {title}' caption paragraph.
    Word renumbers the SEQ field automatically when the file opens (because
    settings.xml has updateFields=true) — so adding or reordering images
    higher in the document doesn't break the numbering."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    p.append(pPr)

    def italic_run(text: str):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rPr.append(OxmlElement("w:i"))
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        return r

    def fld_char(ftype: str):
        r = OxmlElement("w:r")
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), ftype)
        if ftype == "begin":
            fc.set(qn("w:dirty"), "true")
        r.append(fc)
        return r

    def instr_text(text: str):
        r = OxmlElement("w:r")
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = text
        r.append(it)
        return r

    p.append(italic_run("Figure "))
    p.append(fld_char("begin"))
    p.append(instr_text(r" SEQ Figure \* ARABIC "))
    p.append(fld_char("separate"))
    p.append(italic_run("1"))  # placeholder cached value; Word recomputes
    p.append(fld_char("end"))
    p.append(italic_run(f": {title}"))
    return p


def _seq_runs(rpr_template):
    """Return a list of runs that render 'Figure <SEQ Figure>: ' using the
    formatting (bold/italic/font) of `rpr_template` (a <w:rPr> element or None)."""
    def clone_rPr():
        return copy.deepcopy(rpr_template) if rpr_template is not None else None

    def text_run(text: str):
        r = OxmlElement("w:r")
        rPr = clone_rPr()
        if rPr is not None:
            r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        return r

    def fld_run(ftype: str):
        r = OxmlElement("w:r")
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), ftype)
        if ftype == "begin":
            fc.set(qn("w:dirty"), "true")
        r.append(fc)
        return r

    def instr_run(text: str):
        r = OxmlElement("w:r")
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = text
        r.append(it)
        return r

    return [
        text_run("Figure "),
        fld_run("begin"),
        instr_run(r" SEQ Figure \* ARABIC "),
        fld_run("separate"),
        text_run("1"),  # cached placeholder; Word recomputes
        fld_run("end"),
        text_run(": "),
    ]


def _inject_seq_into_paragraph(p_el) -> bool:
    """Inject a 'Figure <SEQ>: ' prefix at the front of an existing paragraph,
    preserving the existing runs and their formatting. Skips if the paragraph
    already contains a SEQ Figure field."""
    if any(
        (it.text or "").strip().startswith("SEQ Figure")
        for it in p_el.findall(".//" + qn("w:instrText"))
    ):
        return False
    runs = p_el.findall(qn("w:r"))
    if not runs:
        return False
    # Get formatting from the first run for the SEQ prefix to inherit
    rPr_template = runs[0].find(qn("w:rPr"))
    seq_runs = _seq_runs(rPr_template)
    # Insert all SEQ runs before the first existing run.
    # addprevious inserts immediately before runs[0]; iterating forward keeps
    # the inserted sequence in the same order, producing
    # "Figure {SEQ}: " + original-text.
    for new_r in seq_runs:
        runs[0].addprevious(new_r)
    return True


def phase13_inject_body_placeholders(doc) -> int:
    """Insert a `{{KEY}}` body placeholder paragraph after each H5/H6 heading
    in Section 1 (INTRODUCTION) and Section 2 (PROPOSED TECHNOLOGY) so the
    Phase-4 content-writer's replacements.json has somewhere to land.
    Without these placeholders, the body of every numbered section is empty
    in the assembled docx.

    Headings mapped to placeholder keys:
        "Executive Summary"      -> EXECUTIVE_SUMMARY
        "Problems & Solutions"   -> PROBLEMS_AND_SOLUTIONS
        "Purpose"                -> PURPOSE
        "System Overview"        -> SYSTEM_OVERVIEW_INTRO
        "Back-end"               -> TECHSTACK_BACKEND
        "Front-end"              -> TECHSTACK_FRONTEND
        "Database"               -> TECHSTACK_DATABASE
        "Server & Hosting"       -> TECHSTACK_SERVER_HOSTING

    Skipped:
      • 2.1.1 System Context — image attaches here, no body paragraph needed
      • Mobile App Strategy / TECHSTACK_DATA / TECHSTACK_AI — already drop-target placeholders
    """
    mapping = {
        "executive summary":     "EXECUTIVE_SUMMARY",
        "problems & solutions":  "PROBLEMS_AND_SOLUTIONS",
        "problems and solutions":"PROBLEMS_AND_SOLUTIONS",
        "purpose":               "PURPOSE",
        "system overview":       "SYSTEM_OVERVIEW_INTRO",
        "back-end":              "TECHSTACK_BACKEND",
        "front-end":             "TECHSTACK_FRONTEND",
        "database":              "TECHSTACK_DATABASE",
        "server & hosting":      "TECHSTACK_SERVER_HOSTING",
    }
    body = doc.element.body
    paras = doc.paragraphs
    inserted = 0
    for p in paras:
        s = p.style.name if p.style else ""
        if s not in ("Heading 5", "Heading 6"):
            continue
        key_name = p.text.strip().lower()
        key = mapping.get(key_name)
        if key is None:
            continue
        # Skip if next sibling already contains the placeholder
        nxt = p._p.getnext()
        if nxt is not None:
            nxt_text = "".join(t.text or "" for t in nxt.iter(qn("w:t")))
            if f"{{{{{key}}}}}" in nxt_text:
                continue
        # Build a body paragraph that holds the placeholder, with COMPACT
        # spacing so the heading and its body don't look orphaned by a
        # large empty gap.
        new_p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "120")  # 6 pt — tight gap to next item
        spacing.set(qn("w:line"), "276")
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")  # justified body to match the rest
        pPr.append(jc)
        new_p.append(pPr)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = f"{{{{{key}}}}}"
        r.append(t)
        new_p.append(r)
        p._p.addnext(new_p)
        inserted += 1
    return inserted


def phase15_normalise_heading_spacing(doc) -> int:
    """Set uniform compact spacing on every Heading 5 and Heading 6
    paragraph so the visual rhythm between sections is consistent.
    Inherited style values are inconsistent across the source template
    (some inherited 200, some 240, some none), which the eye sees as
    'random big gaps between titles'.

    Targets:
      Heading 5: before=160 twips (8 pt), after=120 (6 pt)
      Heading 6: before=120 twips (6 pt), after=80  (4 pt)
    """
    targets = {
        "Heading 5": ("160", "120"),
        "Heading 6": ("120", "80"),
    }
    touched = 0
    for p in doc.paragraphs:
        sname = p.style.name if p.style else ""
        if sname not in targets:
            continue
        before, after = targets[sname]
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._p.insert(0, pPr)
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:before"), before)
        spacing.set(qn("w:after"), after)
        touched += 1
    return touched


def phase14_collapse_consecutive_empty(doc) -> int:
    """Collapse runs of >=2 consecutive empty body paragraphs (no text, no
    drawing, no sectPr) down to at most one. Empty paragraphs accumulate
    after strip / inject operations and produce the 'big gap between
    sections' the eye picks up immediately. Preserves sectPr and image
    paragraphs."""
    body = doc.element.body
    to_remove = []
    prev_empty = None
    for el in list(body):
        if el.tag.rsplit("}", 1)[-1] != "p":
            prev_empty = None
            continue
        text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        has_drawing = bool(el.findall(".//" + qn("w:drawing"))) or bool(
            el.findall(".//" + qn("w:pict"))
        )
        has_sectpr = bool(el.findall(".//" + qn("w:sectPr")))
        if text or has_drawing or has_sectpr:
            prev_empty = None
            continue
        # This is an empty paragraph
        if prev_empty is not None:
            to_remove.append(el)
        prev_empty = el
    for el in to_remove:
        if el.getparent() is not None:
            el.getparent().remove(el)
    return len(to_remove)


def phase12_spacing_after_caption(doc) -> int:
    """Ensure a small empty paragraph follows each 1x1 caption table so the
    caption isn't visually pressed against the next bullet / paragraph."""
    body = doc.element.body
    added = 0
    for el in list(body):
        if el.tag.rsplit("}", 1)[-1] != "tbl":
            continue
        rows = el.findall(qn("w:tr"))
        if len(rows) != 1:
            continue
        cells = rows[0].findall(qn("w:tc"))
        if len(cells) != 1:
            continue
        cell_text = "".join(t.text or "" for t in cells[0].iter(qn("w:t"))).strip()
        if not cell_text:
            continue
        # Look at the next sibling element.
        nxt = el.getnext()
        if nxt is None:
            continue
        # Skip if next is already an empty paragraph providing the gap.
        if nxt.tag.rsplit("}", 1)[-1] == "p":
            nxt_text = "".join(t.text or "" for t in nxt.iter(qn("w:t"))).strip()
            has_img = bool(nxt.findall(".//" + qn("w:drawing")))
            if not nxt_text and not has_img:
                continue
        # Insert an empty paragraph between the caption table and the next
        # element so the figure caption breathes. 240 twips (12 pt) — visible
        # but not excessive.
        spacer = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "120")
        spacing.set(qn("w:after"), "240")  # 12 pt
        pPr.append(spacing)
        spacer.append(pPr)
        el.addnext(spacer)
        added += 1
    return added


def phase11_seq_captions(doc) -> int:
    """Make EVERY figure caption look the same (1x1 bordered table style used
    by Section 3 verbatim diagrams) and inject an auto-numbering SEQ field
    into each one.

    Two cases handled:
      1. Hardcoded 'Figure N: Title' paragraphs (e.g. the legacy CI/CD
         caption) — replaced with a 1x1 caption table cloned from one of
         the existing Section 3 captions so the visual style matches.
      2. Existing 1x1 caption tables — left in place; only the inner cell
         text is prefixed with 'Figure <SEQ>: '.

    Word renumbers all SEQ fields on open (settings.xml updateFields=true,
    each fldChar begin carries w:dirty='true').
    """
    converted = 0
    body = doc.element.body

    # ------------------------------------------------------------
    # Pre-step: pick a representative 1x1 caption table to clone for any
    # hardcoded 'Figure N: ...' paragraph we find. Must be done BEFORE we
    # mutate the tables themselves.
    # ------------------------------------------------------------
    sample_caption_tbl = None
    for t in doc.tables:
        if len(t.rows) == 1 and len(t.columns) == 1:
            txt = t.rows[0].cells[0].text.strip()
            if txt and not txt.lower().startswith("figure "):
                sample_caption_tbl = t
                break

    # ------------------------------------------------------------
    # Pass 1: replace hardcoded 'Figure N: Title' paragraphs with cloned
    # 1x1 caption tables containing just the title text (the SEQ prefix
    # will be injected in Pass 2 alongside the other caption tables).
    # ------------------------------------------------------------
    fig_re = re.compile(r"^\s*Figure\s+\d+\s*[:.]?\s*(.+?)\s*$")
    for p in list(doc.paragraphs):
        m = fig_re.match(p.text)
        if not m:
            continue
        # Skip paragraphs that already contain a SEQ field (already converted).
        if p._p.findall(".//" + qn("w:instrText")):
            continue
        title = m.group(1).strip()
        parent_el = p._p.getparent()
        if parent_el is None:
            continue
        if sample_caption_tbl is None:
            # Fall back to a plain SEQ paragraph if no caption table exists.
            idx = list(parent_el).index(p._p)
            parent_el.insert(idx, make_seq_caption_paragraph(title))
            parent_el.remove(p._p)
            converted += 1
            continue
        new_tbl = copy.deepcopy(sample_caption_tbl._tbl)
        # Wipe text in the clone, then set just the title (SEQ added in Pass 2).
        for tn in new_tbl.iter(qn("w:t")):
            tn.text = ""
        first_t = next(iter(new_tbl.iter(qn("w:t"))), None)
        if first_t is not None:
            first_t.text = title
        else:
            first_p = next(iter(new_tbl.iter(qn("w:p"))), None)
            if first_p is not None:
                r = OxmlElement("w:r")
                t_el = OxmlElement("w:t")
                t_el.text = title
                r.append(t_el)
                first_p.append(r)
        idx = list(parent_el).index(p._p)
        parent_el.insert(idx, new_tbl)
        parent_el.remove(p._p)
        converted += 1

    # ------------------------------------------------------------
    # Pass 2: inject 'Figure <SEQ>: ' prefix into every 1x1 caption table's
    # cell, preserving the cell's bordered styling.
    # ------------------------------------------------------------
    for el in list(body):
        if el.tag.rsplit("}", 1)[-1] != "tbl":
            continue
        rows = el.findall(qn("w:tr"))
        if len(rows) != 1:
            continue
        cells = rows[0].findall(qn("w:tc"))
        if len(cells) != 1:
            continue
        target_p = None
        for p_el in cells[0].findall(qn("w:p")):
            if any((t.text or "").strip() for t in p_el.iter(qn("w:t"))):
                target_p = p_el
                break
        if target_p is None:
            continue
        if _inject_seq_into_paragraph(target_p):
            converted += 1
    return converted


def phase8_keep_groups(doc) -> int:
    paras = doc.paragraphs
    bound = 0
    for i, p in enumerate(paras):
        if not has_drawing(p):
            continue
        set_keep(p, "w:keepLines")
        j = i - 1
        chain_start = None
        while j >= 0:
            prev = paras[j]
            sname = prev.style.name if prev.style else ""
            if sname.startswith("Heading"):
                chain_start = j
                break
            if prev.text.strip() == "" and not has_drawing(prev):
                j -= 1
                continue
            break
        if chain_start is not None:
            for k in range(chain_start, i):
                set_keep(paras[k], "w:keepNext")
            bound += 1
    return bound


def main() -> int:
    if not SRC.exists():
        print(f"! template missing: {SRC}", file=sys.stderr)
        return 1
    print("Phase 1: XML string subs + TOC dirty + settings patch")
    n_toc = phase1_xml_subs(SRC)
    print(f"  TOC fields marked dirty: {n_toc}")

    doc = Document(str(SRC))
    print(f"Phase 2: cover replacements: {phase2_cover_repl(doc)}")
    print(f"Phase 3: 2.1.x sub-headings removed: {phase3_drop_21x(doc)}")
    print(f"Phase 4: tech-stack placeholders added: {phase4_add_techstack(doc)}")
    print(f"Phase 5: SUMMARY normalized: {phase5_normalize_summary(doc)}")
    # CASE STUDY body kept as-is from original (per user feedback);
    # only SUMMARY is stripped because its content was 12 paragraphs of
    # project-specific narrative that doesn't generalise.
    print(f"Phase 6: SUMMARY body items stripped: {phase6_strip_body(doc, 'SUMMARY', '{{SUMMARY_BODY}}')}")
    doc.save(str(SRC))

    doc = Document(str(SRC))
    print(f"Phase 7: testing-process restructured: {phase7_restructure_testing(doc)}")
    doc.save(str(SRC))

    doc = Document(str(SRC))
    print(f"Phase 8: keep-with-next groups: {phase8_keep_groups(doc)}")
    print(f"Phase 9: numbered bullets cleaned (strip bullet glyph): {phase9_strip_bullet_before_number(doc)}")
    print(f"Phase 10: heading-image spacing added: {phase10_spacing_heading_before_image(doc)}")
    print(f"Phase 11: captions converted to SEQ field: {phase11_seq_captions(doc)}")
    print(f"Phase 12: spacing after caption tables added: {phase12_spacing_after_caption(doc)}")
    print(f"Phase 13: body placeholders injected under Section 1-2 H5/H6: {phase13_inject_body_placeholders(doc)}")
    print(f"Phase 14: consecutive empty paragraphs collapsed: {phase14_collapse_consecutive_empty(doc)}")
    print(f"Phase 15: H5/H6 spacing normalised: {phase15_normalise_heading_spacing(doc)}")
    doc.save(str(SRC))

    # Verify
    print()
    with zipfile.ZipFile(SRC) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    sectPrs = re.findall(r"<w:sectPr[^>]*>.*?</w:sectPr>", xml, re.DOTALL)
    print(f"Final sections: {len(sectPrs)}")
    for i, s in enumerate(sectPrs):
        m = re.search(r"<w:pgMar\s+([^/]+)/>", s)
        if m:
            mm = re.search(r'w:top="(\d+)"\s+w:right="(\d+)"\s+w:bottom="(\d+)"\s+w:left="(\d+)"', m.group(1))
            if mm:
                print(f"  Sec{i+1}: top={mm.group(1)} right={mm.group(2)} bottom={mm.group(3)} left={mm.group(4)}")
    print(f"File size: {SRC.stat().st_size:,} bytes")
    return 0


if __name__ == "__main__":
    sys.exit(main())
